#!/usr/bin/env python3
"""
MetaScout - Advanced File Metadata Analysis Tool
------------------------------------------------
A professional-grade CLI tool for extracting, analyzing, and reporting metadata
across multiple file types with security-focused insights and batch processing capabilities.
"""

import argparse
import os
import json
import logging
import hashlib
import datetime
import concurrent.futures
import csv
import tempfile
import shutil
import re
import sys
from typing import Dict, Optional, List, Tuple, Set, Any, Union, Callable
from pathlib import Path
from dataclasses import dataclass, field, asdict
import textwrap
import importlib

# --- Optional Dependencies Management ---
OPTIONAL_DEPENDENCIES = {
    'magic': True,   # We'll use a fallback if not available
    'yara-python': False,
    'pyssdeep': False,
    'docx': False,
    'openpyxl': False,
    'olefile': False,
    'pefile': False,
    'pyelftools': False,
    'macholib': False
}

# Standard library imports
import xml.etree.ElementTree as ET  # For XML parsing

# Required third-party imports - these will raise errors if missing
try:
    from PIL import Image, ExifTags, IptcImagePlugin
    from PyPDF2 import PdfReader
    import tabulate  # For table output formatting
    import mutagen  # Enhanced audio metadata
    from mutagen.id3 import ID3
    from mutagen.mp4 import MP4
    from mutagen.flac import FLAC
    from mutagen.wave import WAVE
    import exifread  # For more robust EXIF reading
    import colorama  # For colored terminal output
    from colorama import Fore, Style
    import tqdm  # For progress bars
    from cryptography.hazmat.primitives import hashes  # For cryptographic hashing
except ImportError as e:
    print(f"Error: Required dependency not found: {e}")
    print("Please install required dependencies with: pip install -r requirements.txt")
    sys.exit(1)

# Optional dependency: python-magic-bin (with fallback to mimetypes)
try:
    import libmagic
    import magic
    # Test if it actually works by calling a function
    test = magic.Magic()
    test.from_buffer(b"test")
    OPTIONAL_DEPENDENCIES['magic'] = True
except (ImportError, AttributeError, TypeError) as e:
    OPTIONAL_DEPENDENCIES['magic'] = False
    import mimetypes
    # Define a magic-like class using mimetypes
    class FallbackMagic:
        def __init__(self, mime=True):
            self.mime = mime
        
        def from_file(self, filename):
            mime_type, encoding = mimetypes.guess_type(filename)
            if self.mime:
                return mime_type or "application/octet-stream"
            else:
                return f"data file ({mime_type or 'unknown'})"
    
    # Create a global instance with similar API to python-magic
    magic = FallbackMagic()

# Optional dependency: Microsoft Office document handling
try:
    import docx  # For MS Word documents
    OPTIONAL_DEPENDENCIES['docx'] = True
except ImportError:
    print("Warning: python-docx not available. Word document analysis will be limited.")

try:
    import openpyxl  # For Excel documents
    OPTIONAL_DEPENDENCIES['openpyxl'] = True
except ImportError:
    print("Warning: openpyxl not available. Excel document analysis will be limited.")

try:
    import olefile  # For OLE files (older Office documents)
    OPTIONAL_DEPENDENCIES['olefile'] = True
except ImportError:
    print("Warning: olefile not available. Legacy Office document analysis will be limited.")

# Optional dependency: Executable analysis
try:
    import pefile  # For PE file analysis
    OPTIONAL_DEPENDENCIES['pefile'] = True
except ImportError:
    print("Warning: pefile not available. Windows executable analysis will be limited.")

try:
    from elftools.elf.elffile import ELFFile  # For ELF file analysis
    OPTIONAL_DEPENDENCIES['pyelftools'] = True
except ImportError:
    print("Warning: pyelftools not available. Linux executable analysis will be limited.")

try:
    from macholib.MachO import MachO  # For Mach-O file analysis
    OPTIONAL_DEPENDENCIES['macholib'] = True
except ImportError:
    print("Warning: macholib not available. macOS executable analysis will be limited.")

# Optional dependency: YARA pattern matching
try:
    import yara  # For pattern matching
    OPTIONAL_DEPENDENCIES['yara'] = True
except FileNotFoundError:
    print("Warning: yara-python not available. YARA rule scanning will be disabled.")

# Optional dependency: Fuzzy hashing
try:
    import pyssdeep  # For fuzzy hashing
    OPTIONAL_DEPENDENCIES['pyssdeep'] = True
except ImportError:
    print("Warning: pyssdeep not available. Fuzzy hash comparison will be disabled.")

# Initialize colorama for cross-platform colored output
colorama.init(autoreset=True)

# --- Constants ---
VERSION = "1.0.0"
SUPPORTED_EXTENSIONS = {
    'images': ['.jpg', '.jpeg', '.png', '.tiff', '.gif', '.bmp', '.webp'],
    'documents': ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt', '.odt', '.ods', '.odp'],
    'audio': ['.mp3', '.wav', '.flac', '.m4a', '.ogg', '.aac'],
    'video': ['.mp4', '.avi', '.mkv', '.mov', '.wmv'],
    'archives': ['.zip', '.rar', '.tar', '.gz', '.7z'],
    'executables': ['.exe', '.dll', '.so', '.dylib'],
    'scripts': ['.js', '.py', '.sh', '.ps1', '.bat']
}

PRIVACY_CONCERNS = {
    'gps_data': re.compile(r'GPS|geotag|location', re.IGNORECASE),
    'email': re.compile(r'([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)'),
    'phone': re.compile(r'\b(?:\+\d{1,2}\s)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}\b'),
    'ssn': re.compile(r'\b\d{3}-\d{2}-\d{4}\b')
}

SECURITY_CONCERNS = {
    'software_names': ['adobe', 'photoshop', 'microsoft', 'office', 'acrobat'],
    'suspicious_entries': ['hidden', 'password', 'encrypted', 'script', 'macro']
}

# --- Data Models ---

@dataclass
class MetadataFinding:
    """Represents a finding or insight from metadata analysis."""
    type: str  # Type of finding (privacy, security, inconsistency.)
    description: str  # Human-readable description
    severity: str  # high, medium, low
    data: Dict = field(default_factory=dict)  # Additional data related to the finding

@dataclass
class FileMetadata:
    """Container for file metadata and analysis results."""
    file_path: str
    file_type: str
    file_size: int
    mime_type: str
    hashes: Dict[str, str] = field(default_factory=dict)
    creation_time: Optional[str] = None
    modification_time: Optional[str] = None
    access_time: Optional[str] = None
    metadata: Dict = field(default_factory=dict)
    findings: List[MetadataFinding] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        """Convert to dictionary for serialization."""
        return asdict(self)

# --- Core System Functions ---

def configure_logging(log_file: str, verbose: bool = False) -> None:
    """Configure logging system with appropriate levels and handlers."""
    log_level = logging.DEBUG if verbose else logging.INFO
    
    # Create a formatter
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    
    # Configure file handler
    file_handler = logging.FileHandler(log_file)
    file_handler.setLevel(logging.DEBUG)  # Always log details to file
    file_handler.setFormatter(formatter)
    
    # Configure console handler with color
    console_handler = logging.StreamHandler()
    console_handler.setLevel(log_level)
    
    # Custom formatter for console with colors
    class ColoredFormatter(logging.Formatter):
        FORMATS = {
            logging.DEBUG: Fore.CYAN + "%(message)s" + Style.RESET_ALL,
            logging.INFO: "%(message)s",
            logging.WARNING: Fore.YELLOW + "%(message)s" + Style.RESET_ALL,
            logging.ERROR: Fore.RED + "%(message)s" + Style.RESET_ALL,
            logging.CRITICAL: Fore.RED + Style.BRIGHT + "%(message)s" + Style.RESET_ALL
        }

        def format(self, record):
            log_fmt = self.FORMATS.get(record.levelno)
            formatter = logging.Formatter(log_fmt)
            return formatter.format(record)
            
    console_handler.setFormatter(ColoredFormatter())
    
    # Get the root logger and add handlers
    logger = logging.getLogger()
    logger.setLevel(logging.DEBUG)  # Allow all logs to be processed
    
    # Remove existing handlers to avoid duplicates on reconfiguration
    for handler in logger.handlers[:]:
        logger.removeHandler(handler)
        
    logger.addHandler(file_handler)
    
    # Only add console handler if not in quiet mode
    if verbose:
        logger.addHandler(console_handler)

def compute_file_hashes(file_path: str) -> Dict[str, str]:
    """Compute multiple secure hashes for a file."""
    hashes = {}
    
    # Define hash algorithms to use
    hash_algorithms = {
        'md5': hashlib.md5(),
        'sha1': hashlib.sha1(),
        'sha256': hashlib.sha256(),
        'sha512': hashlib.sha512()
    }
    
    try:
        with open(file_path, 'rb') as f:
            # Read file in chunks to handle large files efficiently
            for chunk in iter(lambda: f.read(4096), b''):
                for hash_obj in hash_algorithms.values():
                    hash_obj.update(chunk)
        
        # Get hex digests
        for name, hash_obj in hash_algorithms.items():
            hashes[name] = hash_obj.hexdigest()
            
        return hashes
    except Exception as e:
        logging.error(f"Failed to compute hashes for {file_path}: {e}")
        return {'error': str(e)}

def detect_file_type(file_path: str) -> Tuple[str, str]:
    """
    Detect file type using libmagic or fallback and return MIME type and more specific description.
    """
    try:
        if OPTIONAL_DEPENDENCIES['magic']:
            # Using actual python-magic
            mime = magic.Magic(mime=True)
            mime_type = mime.from_file(file_path)
            
            # Get more detailed description
            magic_desc = magic.Magic()
            description = magic_desc.from_file(file_path)
        else:
            # Using our fallback
            mime_type = magic.from_file(file_path)
            description = f"File: {os.path.basename(file_path)}"
            
            # Try to get more info from extension
            ext = os.path.splitext(file_path)[1].lower()
            if ext:
                description += f" ({ext} file)"
        
        return mime_type, description
    except Exception as e:
        logging.error(f"Failed to detect file type for {file_path}: {e}")
        # Last resort - guess from extension
        ext = os.path.splitext(file_path)[1].lower()
        if ext:
            import mimetypes
            mime_type, _ = mimetypes.guess_type(file_path)
            return mime_type or "unknown/unknown", f"Unknown file type with extension {ext}"
        return "unknown/unknown", "Unknown file type"

def get_file_timestamps(file_path: str) -> Dict[str, str]:
    """Get file creation, modification, and access times."""
    try:
        stat_info = os.stat(file_path)
        return {
            'creation_time': datetime.datetime.fromtimestamp(stat_info.st_ctime).isoformat(),
            'modification_time': datetime.datetime.fromtimestamp(stat_info.st_mtime).isoformat(),
            'access_time': datetime.datetime.fromtimestamp(stat_info.st_atime).isoformat()
        }
    except Exception as e:
        logging.error(f"Failed to get timestamps for {file_path}: {e}")
        return {}

# --- Metadata Extraction Functions ---

def extract_image_metadata(file_path: str) -> Dict:
    """Extract comprehensive metadata from image files."""
    metadata = {
        'basic': {},
        'exif': {},
        'iptc': {},
        'xmp': {},
    }
    
    try:
        # Use PIL for basic image info
        with Image.open(file_path) as img:
            metadata['basic'] = {
                'format': img.format,
                'mode': img.mode,
                'width': img.width,
                'height': img.height,
                'color_profile': img.info.get('icc_profile', 'None')
            }
            
            # Extract EXIF data
            exif_data = img._getexif()
            if exif_data:
                # Convert numeric EXIF tags to readable names
                metadata['exif'] = {
                    ExifTags.TAGS.get(tag, tag): str(value)
                    for tag, value in exif_data.items()
                    if tag in ExifTags.TAGS
                }
                
                # Handle GPS info specially
                if 'GPSInfo' in metadata['exif'] and isinstance(metadata['exif']['GPSInfo'], dict):
                    gps_info = {}
                    for key, val in metadata['exif']['GPSInfo'].items():
                        gps_key = ExifTags.GPSTAGS.get(key, key)
                        gps_info[gps_key] = val
                    metadata['exif']['GPSInfo'] = gps_info
            
            # Extract IPTC data
            iptc = IptcImagePlugin.getiptcinfo(img)
            if iptc:
                metadata['iptc'] = {str(key): value.decode('utf-8', 'ignore') 
                                    if isinstance(value, bytes) else str(value) 
                                    for key, value in iptc.items()}
            
            # Extract XMP data if present
            if 'xmp' in img.info:
                # Parse XML structure of XMP
                try:
                    xmp_str = img.info['xmp'].decode('utf-8', 'ignore') if isinstance(img.info['xmp'], bytes) else img.info['xmp']
                    root = ET.fromstring(xmp_str)
                    metadata['xmp'] = {'raw': xmp_str}
                    
                    # Extract all namespaces
                    namespaces = {k: v for k, v in root.attrib.items() if k.startswith('xmlns:')}
                    
                    # Extract key properties from XMP data
                    for elem in root.iter():
                        tag = elem.tag.split('}')[-1]
                        if elem.text and elem.text.strip():
                            metadata['xmp'][tag] = elem.text.strip()
                except Exception as e:
                    logging.warning(f"Failed to parse XMP data: {e}")
                    metadata['xmp'] = {'raw': str(img.info['xmp'])}
        
        # Use ExifRead for more thorough EXIF extraction
        with open(file_path, 'rb') as f:
            exif_tags = exifread.process_file(f, details=True)
            if exif_tags:
                # Merge with existing EXIF data, preferring ExifRead values
                for tag, value in exif_tags.items():
                    tag_name = tag.replace(' ', '_')
                    metadata['exif'][tag_name] = str(value)
        
        return metadata
    except Exception as e:
        logging.error(f"Error extracting image metadata from {file_path}: {e}")
        return {'error': str(e)}

def extract_pdf_metadata(file_path: str) -> Dict:
    """Extract metadata from PDF files."""
    metadata = {
        'document_info': {},
        'xmp_metadata': {},
        'page_info': {},
        'security': {}
    }
    
    try:
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            
            # Basic document info
            if reader.metadata:
                for key, value in reader.metadata.items():
                    if isinstance(key, str):
                        clean_key = key.lstrip('/')
                        metadata['document_info'][clean_key] = str(value)
            
            # Page information
            metadata['page_info'] = {
                'page_count': len(reader.pages),
                'pages': []
            }
            
            # Sample the first few pages for detailed info
            max_pages_to_sample = min(len(reader.pages), 5)
            for i in range(max_pages_to_sample):
                page = reader.pages[i]
                page_info = {
                    'index': i,
                    'rotation': page.get('/Rotate', 0),
                    'size': {'width': page.mediabox.width, 'height': page.mediabox.height}
                }
                metadata['page_info']['pages'].append(page_info)
            
            # Security information
            metadata['security'] = {
                'encrypted': reader.is_encrypted,
                'permissions': {
                    'printing': not reader.is_encrypted
                }
            }
            
            if reader.is_encrypted:
                metadata['security']['permissions'] = {
                    'printing': reader.can_print,
                    'modification': reader.can_modify,
                    'copy': reader.can_copy,
                    'annotation': reader.can_annotate,
                    'filling_forms': reader.can_fill_forms,
                    'extract_content': reader.can_extract,
                    'assemble_doc': reader.can_assemble
                }
            
            return metadata
    except Exception as e:
        logging.error(f"Error extracting PDF metadata from {file_path}: {e}")
        return {'error': str(e)}

def extract_office_document_metadata(file_path: str) -> Dict:
    """Extract metadata from Office documents (DOCX, XLSX, etc.)."""
    metadata = {
        'document_properties': {},
        'custom_properties': {},
        'content_stats': {}
    }
    
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext in ['.docx', '.doc']:
            if ext == '.docx':
                # Modern DOCX files
                doc = docx.Document(file_path)
                
                # Core properties
                core_props = doc.core_properties
                metadata['document_properties'] = {
                    'author': core_props.author,
                    'created': str(core_props.created) if core_props.created else None,
                    'last_modified_by': core_props.last_modified_by,
                    'modified': str(core_props.modified) if core_props.modified else None,
                    'title': core_props.title,
                    'subject': core_props.subject,
                    'keywords': core_props.keywords,
                    'language': core_props.language,
                    'category': core_props.category,
                    'version': core_props.revision
                }
                
                # Content statistics
                metadata['content_stats'] = {
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                    'page_count': 'unknown',  # Requires rendering to determine
                    'word_count': sum(len(p.text.split()) for p in doc.paragraphs),
                    'character_count': sum(len(p.text) for p in doc.paragraphs)
                }
            
            elif ext == '.doc':
                # Legacy DOC files
                if olefile.isOleFile(file_path):
                    ole = olefile.OleFile(file_path)
                    
                    # Extract SummaryInformation stream
                    if ole.exists('\\x05SummaryInformation'):
                        si_stream = ole.openstream('\\x05SummaryInformation')
                        from oletools.oleid import OleID
                        oledata = OleID(file_path)
                        indicators = oledata.check()
                        
                        for indicator in indicators:
                            if indicator.name == 'author':
                                metadata['document_properties']['author'] = indicator.value
                            if indicator.name == 'creation_time':
                                metadata['document_properties']['created'] = str(indicator.value)
                            if indicator.name == 'last_saved_time':
                                metadata['document_properties']['modified'] = str(indicator.value)
                            if indicator.name == 'vba_macros':
                                metadata['security'] = {'macros_present': indicator.value}
                    
                    ole.close()
        
        elif ext in ['.xlsx', '.xls']:
            if ext == '.xlsx':
                # Modern XLSX files
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                
                # Basic properties
                metadata['document_properties'] = {
                    'creator': wb.properties.creator,
                    'created': str(wb.properties.created) if wb.properties.created else None,
                    'last_modified_by': wb.properties.lastModifiedBy,
                    'modified': str(wb.properties.modified) if wb.properties.modified else None,
                    'title': wb.properties.title,
                    'subject': wb.properties.subject,
                    'keywords': wb.properties.keywords,
                    'category': wb.properties.category
                }
                
                # Content statistics
                sheet_stats = []
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    sheet_info = {
                        'name': sheet_name,
                        'dimensions': sheet.calculate_dimension(),
                        'protection': sheet.protection.sheet
                    }
                    sheet_stats.append(sheet_info)
                
                metadata['content_stats'] = {
                    'sheet_count': len(wb.sheetnames),
                    'sheets': sheet_stats
                }
                
                wb.close()
            
            elif ext == '.xls':
                # Legacy XLS files
                if olefile.isOleFile(file_path):
                    ole = olefile.OleFile(file_path)
                    
                    # Extract SummaryInformation stream
                    if ole.exists('\\x05SummaryInformation'):
                        # Same approach as DOC files
                        si_stream = ole.openstream('\\x05SummaryInformation')
                        from oletools.oleid import OleID
                        oledata = OleID(file_path)
                        indicators = oledata.check()
                        
                        for indicator in indicators:
                            if indicator.name == 'author':
                                metadata['document_properties']['author'] = indicator.value
                            if indicator.name == 'creation_time':
                                metadata['document_properties']['created'] = str(indicator.value)
                            if indicator.name == 'last_saved_time':
                                metadata['document_properties']['modified'] = str(indicator.value)
                            if indicator.name == 'vba_macros':
                                metadata['security'] = {'macros_present': indicator.value}
                    
                    ole.close()
        
        return metadata
    except Exception as e:
        logging.error(f"Error extracting Office document metadata from {file_path}: {e}")
        return {'error': str(e)}

def extract_audio_metadata(file_path: str) -> Dict:
    """Extract metadata from audio files (MP3, FLAC, WAV, etc.)."""
    metadata = {
        'audio_properties': {},
        'tags': {}
    }
    
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        # General approach using mutagen
        audio = mutagen.File(file_path)
        
        if audio:
            # Audio properties
            if hasattr(audio.info, 'length'):
                metadata['audio_properties']['duration'] = audio.info.length
            if hasattr(audio.info, 'bitrate'):
                metadata['audio_properties']['bitrate'] = audio.info.bitrate
            if hasattr(audio.info, 'sample_rate'):
                metadata['audio_properties']['sample_rate'] = audio.info.sample_rate
            if hasattr(audio.info, 'channels'):
                metadata['audio_properties']['channels'] = audio.info.channels
            
            # Tags
            for key, value in audio.items():
                if isinstance(value, list) and len(value) == 1:
                    metadata['tags'][key] = str(value[0])
                else:
                    metadata['tags'][key] = str(value)
        
        # Format-specific handling
        if ext == '.mp3':
            try:
                id3 = ID3(file_path)
                for frame in id3.values():
                    frame_name = frame.__class__.__name__
                    if frame_name not in metadata['tags']:
                        metadata['tags'][frame_name] = str(frame)
            except Exception as e:
                logging.warning(f"Failed to extract ID3 data: {e}")
        
        elif ext == '.flac':
            try:
                flac = FLAC(file_path)
                for key, value in flac.items():
                    if key not in metadata['tags']:
                        metadata['tags'][key] = value[0] if isinstance(value, list) and value else value
            except Exception as e:
                logging.warning(f"Failed to extract FLAC data: {e}")
        
        elif ext == '.wav':
            try:
                wav = WAVE(file_path)
                for key, value in wav.items():
                    if key not in metadata['tags']:
                        metadata['tags'][key] = value[0] if isinstance(value, list) and value else value
            except Exception as e:
                logging.warning(f"Failed to extract WAV data: {e}")
        
        return metadata
    except Exception as e:
        logging.error(f"Error extracting audio metadata from {file_path}: {e}")
        return {'error': str(e)}

def extract_video_metadata(file_path: str) -> Dict:
    """Extract metadata from video files."""
    metadata = {
        'video_properties': {},
        'audio_streams': [],
        'subtitle_streams': []
    }
    
    try:
        # Try to use FFProbe if available (requires ffmpeg installation)
        import subprocess
        
        # Check if ffprobe is available
        try:
            subprocess.run(['ffprobe', '-version'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
            has_ffprobe = True
        except (subprocess.SubprocessError, FileNotFoundError):
            has_ffprobe = False
        
        if has_ffprobe:
            # Get video stream info
            cmd = [
                'ffprobe', 
                '-v', 'quiet', 
                '-print_format', 'json', 
                '-show_format', 
                '-show_streams', 
                file_path
            ]
            
            result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
            output = json.loads(result.stdout)
            
            # Format information
            if 'format' in output:
                metadata['video_properties'] = {
                    'format': output['format'].get('format_name', 'unknown'),
                    'duration': float(output['format'].get('duration', 0)),
                    'size': int(output['format'].get('size', 0)),
                    'bit_rate': int(output['format'].get('bit_rate', 0)) if 'bit_rate' in output['format'] else 0
                }
                
                # Additional format tags
                if 'tags' in output['format']:
                    for tag, value in output['format']['tags'].items():
                        metadata['video_properties'][tag.lower()] = value
            
            # Stream information
            if 'streams' in output:
                for stream in output['streams']:
                    stream_type = stream.get('codec_type', 'unknown')
                    
                    if stream_type == 'video':
                        metadata['video_properties'].update({
                            'codec': stream.get('codec_name', 'unknown'),
                            'width': stream.get('width', 0),
                            'height': stream.get('height', 0),
                            'fps': eval(stream.get('r_frame_rate', '0/1')) if '/' in stream.get('r_frame_rate', '0/1') else 0,
                            'pix_fmt': stream.get('pix_fmt', 'unknown')
                        })
                    
                    elif stream_type == 'audio':
                        audio_stream = {
                            'codec': stream.get('codec_name', 'unknown'),
                            'sample_rate': stream.get('sample_rate', 'unknown'),
                            'channels': stream.get('channels', 0),
                            'bit_rate': stream.get('bit_rate', 'unknown') if 'bit_rate' in stream else 'unknown'
                        }
                        metadata['audio_streams'].append(audio_stream)
                    
                    elif stream_type == 'subtitle':
                        subtitle_stream = {
                            'codec': stream.get('codec_name', 'unknown'),
                            'language': stream.get('tags', {}).get('language', 'unknown')
                        }
                        metadata['subtitle_streams'].append(subtitle_stream)
        else:
            # Fallback to basic file information if ffprobe is not available
            metadata['video_properties'] = {
                'note': 'Limited metadata available without ffprobe',
                'file_size': os.path.getsize(file_path)
            }
            
            # Try to use mediainfo if available
            try:
                subprocess.run(['mediainfo', '--version'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
                has_mediainfo = True
            except (subprocess.SubprocessError, FileNotFoundError):
                has_mediainfo = False
            
            if has_mediainfo:
                cmd = ['mediainfo', '--Output=JSON', file_path]
                result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
                output = json.loads(result.stdout)
                
                if 'media' in output and 'track' in output['media']:
                    for track in output['media']['track']:
                        track_type = track.get('@type', '').lower()
                        
                        if track_type == 'video':
                            metadata['video_properties'].update({
                                'codec': track.get('Format', 'unknown'),
                                'width': int(track.get('Width', 0)) if 'Width' in track else 0,
                                'height': int(track.get('Height', 0)) if 'Height' in track else 0,
                                'duration': float(track.get('Duration', 0)) if 'Duration' in track else 0
                            })
        
        return metadata
    except Exception as e:
        logging.error(f"Error extracting video metadata from {file_path}: {e}")
        return {'error': str(e)}

def extract_executable_metadata(file_path: str) -> Dict:
    """Extract metadata from executable files."""
    metadata = {
        'file_headers': {},
        'libraries': [],
        'signatures': {},
    }
    
    try:
        # Basic file signature analysis
        with open(file_path, 'rb') as f:
            header = f.read(16)  # Read first 16 bytes
            
            # Check for PE files (Windows executables)
            if header[0:2] == b'MZ':
                metadata['file_headers']['type'] = 'Windows PE'
                
                # Try to use pefile if available
                try:
                    import pefile
                    pe = pefile.PE(file_path)
                    
                    # Get timestamp
                    timestamp = pe.FILE_HEADER.TimeDateStamp
                    metadata['file_headers']['compile_time'] = datetime.datetime.fromtimestamp(timestamp).isoformat()
                    
                    # Get machine type
                    machine_types = {
                        0x014c: 'x86 (32-bit)',
                        0x0200: 'IA64',
                        0x8664: 'x64 (AMD64)'
                    }
                    metadata['file_headers']['machine'] = machine_types.get(pe.FILE_HEADER.Machine, f'Unknown ({pe.FILE_HEADER.Machine:04X})')
                    
                    # Get imported DLLs
                    if hasattr(pe, 'DIRECTORY_ENTRY_IMPORT'):
                        for entry in pe.DIRECTORY_ENTRY_IMPORT:
                            dll_name = entry.dll.decode('utf-8', 'ignore') if entry.dll else 'Unknown'
                            imports = []
                            
                            if hasattr(entry, 'imports'):
                                for imp in entry.imports:
                                    if imp.name:
                                        imports.append(imp.name.decode('utf-8', 'ignore'))
                            
                            metadata['libraries'].append({
                                'name': dll_name,
                                'imports': imports[:5]  # Limit to first 5 for brevity
                            })
                    
                    # Get sections
                    metadata['file_headers']['sections'] = []
                    for section in pe.sections:
                        section_name = section.Name.decode('utf-8', 'ignore').strip('\x00')
                        metadata['file_headers']['sections'].append({
                            'name': section_name,
                            'size': section.SizeOfRawData,
                            'entropy': section.get_entropy()
                        })
                    
                    # Get resources if available
                    if hasattr(pe, 'DIRECTORY_ENTRY_RESOURCE'):
                        metadata['resources'] = []
                        for resource_type in pe.DIRECTORY_ENTRY_RESOURCE.entries:
                            try:
                                resource_type_name = pefile.RESOURCE_TYPE.get(resource_type.id, str(resource_type.id))
                                metadata['resources'].append(resource_type_name)
                            except Exception as e:
                                logging.debug(f"Error parsing resource: {e}")
                
                except ImportError:
                    metadata['note'] = "Install pefile package for detailed PE analysis"
                except Exception as e:
                    logging.error(f"Error parsing PE file: {e}")
            
            # Check for ELF files (Linux executables)
            elif header[0:4] == b'\x7fELF':
                metadata['file_headers']['type'] = 'Linux ELF'
                
                # Try to use pyelftools if available
                try:
                    from elftools.elf.elffile import ELFFile
                    
                    with open(file_path, 'rb') as f:
                        elf = ELFFile(f)
                        
                        # Get basic header info
                        metadata['file_headers']['class'] = elf.elfclass
                        metadata['file_headers']['data_encoding'] = elf.elfdata
                        metadata['file_headers']['machine_type'] = elf.header['e_machine']
                        
                        # Get sections
                        metadata['file_headers']['sections'] = []
                        for section in elf.iter_sections():
                            metadata['file_headers']['sections'].append({
                                'name': section.name,
                                'type': section['sh_type'],
                                'size': section['sh_size']
                            })
                        
                        # Get dynamic symbols if available
                        dynsym = elf.get_section_by_name('.dynsym')
                        if dynsym:
                            for symbol in dynsym.iter_symbols():
                                if symbol.name:
                                    metadata['libraries'].append({
                                        'name': symbol.name,
                                        'type': symbol['st_info']['type']
                                    })
                
                except ImportError:
                    metadata['note'] = "Install pyelftools package for detailed ELF analysis"
                except Exception as e:
                    logging.error(f"Error parsing ELF file: {e}")
            
            # Check for Mach-O files (macOS executables)
            elif header[0:4] in (b'\xfe\xed\xfa\xce', b'\xfe\xed\xfa\xcf', b'\xca\xfe\xba\xbe'):
                metadata['file_headers']['type'] = 'macOS Mach-O'
                
                # Try to use macholib if available
                try:
                    from macholib.MachO import MachO
                    
                    macho = MachO(file_path)
                    metadata['file_headers']['headers'] = []
                    
                    for header in macho.headers:
                        header_info = {
                            'magic': f"0x{header.MH_MAGIC:08x}",
                            'cpu_type': header.header.cputype,
                            'cpu_subtype': header.header.cpusubtype
                        }
                        metadata['file_headers']['headers'].append(header_info)
                    
                    # Get load commands
                    metadata['libraries'] = []
                    for header in macho.headers:
                        for cmd in header.commands:
                            if cmd[0].cmd == 0x0C:  # LC_LOAD_DYLIB
                                if hasattr(cmd[1], 'name'):
                                    lib_name = cmd[1].name.decode('utf-8', 'ignore') if isinstance(cmd[1].name, bytes) else cmd[1].name
                                    metadata['libraries'].append({'name': lib_name})
                
                except ImportError:
                    metadata['note'] = "Install macholib package for detailed Mach-O analysis"
                except Exception as e:
                    logging.error(f"Error parsing Mach-O file: {e}")
        
        # Signature verification
        try:
            # For Windows executables, check Authenticode signature
            if sys.platform == 'win32' and os.path.splitext(file_path)[1].lower() in ('.exe', '.dll'):
                import win32api
                import win32security
                
                try:
                    info = win32api.GetFileVersionInfo(file_path, '\\')
                    ms = info['FileVersionMS']
                    ls = info['FileVersionLS']
                    metadata['version'] = f"{ms >> 16}.{ms & 0xFFFF}.{ls >> 16}.{ls & 0xFFFF}"
                    
                    # Get signature info
                    cert_store = win32security.CryptQueryObject(0, file_path, 0x02, 0, 0)
                    if cert_store:
                        metadata['signatures']['signed'] = True
                        metadata['signatures']['verified'] = True  # Simplification, actual verification is more complex
                except:
                    metadata['signatures']['signed'] = False
        except ImportError:
            pass
        
        return metadata
    except Exception as e:
        logging.error(f"Error extracting executable metadata from {file_path}: {e}")
        return {'error': str(e)}

def extract_generic_metadata(file_path: str) -> Dict:
    """Extract basic metadata for any file type."""
    try:
        metadata = {}
        
        # Use python-magic to identify file type
        mime_type, description = detect_file_type(file_path)
        metadata['mime_type'] = mime_type
        metadata['description'] = description
        
        # Get file size
        metadata['size'] = os.path.getsize(file_path)
        
        # Get file timestamps
        timestamps = get_file_timestamps(file_path)
        for key, value in timestamps.items():
            metadata[key] = value
        
        # Get file permissions
        try:
            metadata['permissions'] = oct(os.stat(file_path).st_mode)[-3:]
        except:
            pass
        
        # Compute file hashes
        metadata['hashes'] = compute_file_hashes(file_path)
        
        return metadata
    except Exception as e:
        logging.error(f"Error extracting generic metadata from {file_path}: {e}")
        return {'error': str(e)}

# --- Metadata Analysis Functions ---

def analyze_image_metadata(metadata: Dict) -> List[MetadataFinding]:
    """Analyze image metadata for privacy and security concerns."""
    findings = []
    
    # Check for GPS information (privacy concern)
    if 'exif' in metadata:
        gps_keys = ('GPSInfo', 'GPS')
        for key in gps_keys:
            if key in metadata['exif']:
                findings.append(MetadataFinding(
                    type="privacy",
                    description="GPS location data found in EXIF metadata",
                    severity="high",
                    data={"source": "EXIF", "field": key}
                ))
    
    # Check for camera/device information (potential fingerprinting)
    device_info = {}
    for section in ('exif', 'basic'):
        if section in metadata:
            for key in ('Make', 'Model', 'Software', 'ProcessingSoftware'):
                if key in metadata[section]:
                    device_info[key] = metadata[section][key]
    
    if device_info:
        findings.append(MetadataFinding(
            type="privacy",
            description="Device/software information found",
            severity="medium",
            data={"device_info": device_info}
        ))
    
    # Check for creation/modification dates
    dates = {}
    date_keys = ('DateTime', 'DateTimeOriginal', 'DateTimeDigitized', 'ModifyDate')
    for section in ('exif', 'iptc'):
        if section in metadata:
            for key in date_keys:
                if key in metadata[section]:
                    dates[key] = metadata[section][key]
    
    if dates:
        findings.append(MetadataFinding(
            type="information",
            description="Image creation/modification timestamps found",
            severity="low",
            data={"dates": dates}
        ))
    
    # Check for IPTC contact info (privacy)
    contact_info = {}
    if 'iptc' in metadata:
        contact_fields = ('By-line', 'Credit', 'Source', 'Writer-Editor', 'Contact')
        for field in contact_fields:
            if field in metadata['iptc']:
                contact_info[field] = metadata['iptc'][field]
    
    if contact_info:
        findings.append(MetadataFinding(
            type="privacy",
            description="Creator/contact information found in IPTC data",
            severity="medium",
            data={"contact_info": contact_info}
        ))
    
    # Check for editing software (authenticity)
    editing_software = None
    for section in ('exif', 'xmp'):
        if section in metadata:
            for key in ('Software', 'ProcessingSoftware', 'CreatorTool'):
                if key in metadata[section]:
                    editing_software = metadata[section][key]
                    break
    
    if editing_software and any(sw.lower() in editing_software.lower() for sw in ('photoshop', 'gimp', 'lightroom', 'affinity')):
        findings.append(MetadataFinding(
            type="authenticity",
            description=f"Image edited with {editing_software}",
            severity="medium",
            data={"software": editing_software}
        ))
    
    # Check for XMP history (authenticity)
    if 'xmp' in metadata and 'raw' in metadata['xmp']:
        if 'xmpMM:History' in metadata['xmp']['raw']:
            findings.append(MetadataFinding(
                type="authenticity",
                description="XMP edit history found",
                severity="medium",
                data={"edit_history": "XMP history entries present"}
            ))
    
    # Look for embedded color profiles (potential fingerprinting)
    if 'basic' in metadata and 'color_profile' in metadata['basic'] and metadata['basic']['color_profile'] != 'None':
        findings.append(MetadataFinding(
            type="information",
            description="Embedded color profile found",
            severity="low",
            data={"color_profile": "Present"}
        ))
    
    return findings

def analyze_document_metadata(metadata: Dict) -> List[MetadataFinding]:
    """Analyze document metadata for privacy and security concerns."""
    findings = []
    
    # Check for author information
    author_fields = {}
    
    if 'document_info' in metadata:
        for key in ('Author', 'Creator', 'Producer', 'LastModifiedBy'):
            if key in metadata['document_info']:
                author_fields[key] = metadata['document_info'][key]
    
    if author_fields:
        findings.append(MetadataFinding(
            type="privacy",
            description="Author information found",
            severity="medium",
            data={"author_fields": author_fields}
        ))
    
    # Check for creation software
    creation_software = None
    if 'document_info' in metadata:
        for key in ('Creator', 'Producer'):
            if key in metadata['document_info'] and metadata['document_info'][key]:
                creation_software = metadata['document_info'][key]
                break
    
    if creation_software:
        findings.append(MetadataFinding(
            type="information",
            description=f"Document created with {creation_software}",
            severity="low",
            data={"software": creation_software}
        ))
    
    # Check for security features
    if 'security' in metadata:
        if metadata['security'].get('encrypted', False):
            findings.append(MetadataFinding(
                type="security",
                description="Document is encrypted/password-protected",
                severity="medium",
                data={"encryption": True}
            ))
        
        permissions = metadata['security'].get('permissions', {})
        restricted_permissions = {k: v for k, v in permissions.items() if not v}
        if restricted_permissions:
            findings.append(MetadataFinding(
                type="security",
                description="Document has restricted permissions",
                severity="low",
                data={"restricted_permissions": list(restricted_permissions.keys())}
            ))
    
    # Check for macros in Office documents
    if 'security' in metadata and 'macros_present' in metadata['security'] and metadata['security']['macros_present']:
        findings.append(MetadataFinding(
            type="security",
            description="Document contains macros",
            severity="high",
            data={"macros": True}
        ))
    
    # Check for date inconsistencies
    dates = {}
    if 'document_info' in metadata:
        for key in ('CreationDate', 'ModDate', 'created', 'modified'):
            if key in metadata['document_info'] and metadata['document_info'][key]:
                dates[key] = metadata['document_info'][key]
    
    if len(dates) > 1:
        findings.append(MetadataFinding(
            type="information",
            description="Document timestamp information",
            severity="low",
            data={"dates": dates}
        ))
    
    return findings

def analyze_executable_metadata(metadata: Dict) -> List[MetadataFinding]:
    """Analyze executable metadata for security concerns."""
    findings = []
    
    # Check file type
    if 'file_headers' in metadata and 'type' in metadata['file_headers']:
        file_type = metadata['file_headers']['type']
        findings.append(MetadataFinding(
            type="information",
            description=f"Executable type: {file_type}",
            severity="low",
            data={"executable_type": file_type}
        ))
    
    # Check for high entropy sections (possible encryption/packing)
    high_entropy_sections = []
    if 'file_headers' in metadata and 'sections' in metadata['file_headers']:
        for section in metadata['file_headers']['sections']:
            if isinstance(section, dict) and 'entropy' in section and section['entropy'] > 7.0:
                high_entropy_sections.append(section['name'])
    
    if high_entropy_sections:
        findings.append(MetadataFinding(
            type="security",
            description="High entropy sections detected (possible packing/encryption)",
            severity="high",
            data={"sections": high_entropy_sections}
        ))
    
    # Check for suspicious libraries
    suspicious_libs = []
    if 'libraries' in metadata:
        suspicious_names = ['inject', 'hook', 'crypt', 'keylog', 'screen', 'exploit']
        for lib in metadata['libraries']:
            if isinstance(lib, dict) and 'name' in lib:
                lib_name = lib['name'].lower()
                if any(sus in lib_name for sus in suspicious_names):
                    suspicious_libs.append(lib['name'])
    
    if suspicious_libs:
        findings.append(MetadataFinding(
            type="security",
            description="Potentially suspicious libraries detected",
            severity="high",
            data={"libraries": suspicious_libs}
        ))
    
    # Check digital signature status
    if 'signatures' in metadata:
        if metadata['signatures'].get('signed', False):
            verified = metadata['signatures'].get('verified', False)
            if verified:
                findings.append(MetadataFinding(
                    type="security",
                    description="Executable is digitally signed and verified",
                    severity="low",
                    data={"signature": "verified"}
                ))
            else:
                findings.append(MetadataFinding(
                    type="security",
                    description="Executable is signed but not verified",
                    severity="medium",
                    data={"signature": "unverified"}
                ))
        else:
            findings.append(MetadataFinding(
                type="security",
                description="Executable is not digitally signed",
                severity="medium",
                data={"signature": "unsigned"}
            ))
    
    # Check compilation timestamp
    if 'file_headers' in metadata and 'compile_time' in metadata['file_headers']:
        compile_time = metadata['file_headers']['compile_time']
        findings.append(MetadataFinding(
            type="information",
            description=f"Compilation timestamp: {compile_time}",
            severity="low",
            data={"compile_time": compile_time}
        ))
    
    return findings

def analyze_audio_metadata(metadata: Dict) -> List[MetadataFinding]:
    """Analyze audio metadata for privacy concerns."""
    findings = []
    
    # Check for personal information in tags
    personal_info = {}
    if 'tags' in metadata:
        personal_fields = ['artist', 'composer', 'albumartist', 'author', 'encoder']
        for field in personal_fields:
            for key in metadata['tags']:
                if field.lower() in key.lower() and metadata['tags'][key]:
                    personal_info[key] = metadata['tags'][key]
    
    if personal_info:
        findings.append(MetadataFinding(
            type="privacy",
            description="Personal information found in audio tags",
            severity="medium",
            data={"personal_info": personal_info}
        ))
    
    # Check for geolocation data
    geo_info = {}
    if 'tags' in metadata:
        geo_fields = ['geotag', 'location', 'latitude', 'longitude', 'geo']
        for key in metadata['tags']:
            if any(field in key.lower() for field in geo_fields):
                geo_info[key] = metadata['tags'][key]
    
    if geo_info:
        findings.append(MetadataFinding(
            type="privacy",
            description="Geolocation information found in audio tags",
            severity="high",
            data={"geo_info": geo_info}
        ))
    
    # Check for recording device/software info
    device_info = {}
    if 'tags' in metadata:
        device_fields = ['encoder', 'encodedby', 'encoding', 'source', 'device']
        for key in metadata['tags']:
            if any(field in key.lower() for field in device_fields):
                device_info[key] = metadata['tags'][key]
    
    if device_info:
        findings.append(MetadataFinding(
            type="information",
            description="Recording device/software information found",
            severity="low",
            data={"device_info": device_info}
        ))
    
    # Check for unusual audio properties
    if 'audio_properties' in metadata:
        props = metadata['audio_properties']
        if 'channels' in props and props['channels'] > 2:
            findings.append(MetadataFinding(
                type="information",
                description=f"Unusual audio channel configuration: {props['channels']} channels",
                severity="low",
                data={"channels": props['channels']}
            ))
    
    return findings

def analyze_video_metadata(metadata: Dict) -> List[MetadataFinding]:
    """Analyze video metadata for privacy and security concerns."""
    findings = []
    
    # Check for unusual or high-quality encoding (potential sensitive content)
    if 'video_properties' in metadata:
        props = metadata['video_properties']
        
        # Check for high resolution
        if 'width' in props and 'height' in props:
            if props['width'] >= 3840 and props['height'] >= 2160:
                findings.append(MetadataFinding(
                    type="information",
                    description="4K or higher resolution video",
                    severity="low",
                    data={"resolution": f"{props['width']}x{props['height']}"}
                ))
        
        # Check for unusual aspect ratios
        if 'width' in props and 'height' in props and props['width'] > 0 and props['height'] > 0:
            aspect_ratio = props['width'] / props['height']
            if aspect_ratio < 0.5 or aspect_ratio > 3.0:
                findings.append(MetadataFinding(
                    type="information",
                    description=f"Unusual aspect ratio: {aspect_ratio:.2f}",
                    severity="low",
                    data={"aspect_ratio": aspect_ratio}
                ))
        
        # Check for creation date/time
        for key in props:
            if any(date_key in key.lower() for date_key in ['date', 'time', 'created']):
                findings.append(MetadataFinding(
                    type="information",
                    description="Video creation timestamp found",
                    severity="low",
                    data={"timestamp": {key: props[key]}}
                ))
    
    # Check for geolocation data
    geo_info = {}
    if 'video_properties' in metadata:
        for key, value in metadata['video_properties'].items():
            if any(geo_key in key.lower() for geo_key in ['geo', 'gps', 'location', 'latitude', 'longitude']):
                geo_info[key] = value
    
    if geo_info:
        findings.append(MetadataFinding(
            type="privacy",
            description="Geolocation information found in video metadata",
            severity="high",
            data={"geo_info": geo_info}
        ))
    
    # Check for device information
    device_info = {}
    if 'video_properties' in metadata:
        for key, value in metadata['video_properties'].items():
            if any(dev_key in key.lower() for dev_key in ['device', 'camera', 'make', 'model']):
                device_info[key] = value
    
    if device_info:
        findings.append(MetadataFinding(
            type="privacy",
            description="Recording device information found",
            severity="medium",
            data={"device_info": device_info}
        ))
    
    # Check for multiple audio streams (potential hidden content)
    if 'audio_streams' in metadata and len(metadata['audio_streams']) > 1:
        findings.append(MetadataFinding(
            type="information",
            description=f"Multiple audio streams: {len(metadata['audio_streams'])}",
            severity="low",
            data={"audio_streams_count": len(metadata['audio_streams'])}
        ))
    
    # Check for subtitle streams (potential sensitive content)
    if 'subtitle_streams' in metadata and metadata['subtitle_streams']:
        subtitles = [
            f"{s.get('language', 'unknown')}: {s.get('codec', 'unknown')}" 
            for s in metadata['subtitle_streams']
        ]
        findings.append(MetadataFinding(
            type="information",
            description=f"Subtitle streams found: {len(metadata['subtitle_streams'])}",
            severity="low",
            data={"subtitles": subtitles}
        ))
    
    return findings

def analyze_metadata_generic(metadata: Dict) -> List[MetadataFinding]:
    """Analyze generic metadata applicable to any file type."""
    findings = []
    
    # Check for timestamp anomalies
    timestamps = {}
    for key in ('creation_time', 'modification_time', 'access_time'):
        if key in metadata and metadata[key]:
            timestamps[key] = metadata[key]
    
    if len(timestamps) >= 2:
        # Check for files created after they were modified (suspicious)
        if 'creation_time' in timestamps and 'modification_time' in timestamps:
            creation = datetime.datetime.fromisoformat(timestamps['creation_time'])
            modification = datetime.datetime.fromisoformat(timestamps['modification_time'])
            if creation > modification:
                findings.append(MetadataFinding(
                    type="consistency",
                    description="File creation time is after modification time (suspicious)",
                    severity="medium",
                    data={"timestamps": timestamps}
                ))
    
    # Check MIME type consistency with extension
    if 'mime_type' in metadata:
        mime_type = metadata['mime_type']
        ext = os.path.splitext(metadata.get('file_path', ''))[1].lower()
        
        # Define some common MIME type and extension mappings
        mime_ext_map = {
            'image/jpeg': ['.jpg', '.jpeg'],
            'image/png': ['.png'],
            'application/pdf': ['.pdf'],
            'text/plain': ['.txt', '.text'],
            'application/msword': ['.doc'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx'],
            'audio/mpeg': ['.mp3'],
            'video/mp4': ['.mp4']
        }
        
        if mime_type in mime_ext_map and ext not in mime_ext_map[mime_type]:
            expected_exts = ', '.join(mime_ext_map[mime_type])
            findings.append(MetadataFinding(
                type="consistency",
                description=f"File extension '{ext}' does not match MIME type '{mime_type}' (expected: {expected_exts})",
                severity="medium",
                data={"mime_type": mime_type, "extension": ext, "expected_extensions": mime_ext_map[mime_type]}
            ))
    
    # Check for unusually small or large file sizes
    if 'size' in metadata:
        size = metadata['size']
        if size == 0:
            findings.append(MetadataFinding(
                type="consistency",
                description="File size is zero bytes (empty file)",
                severity="medium",
                data={"size": size}
            ))
        elif size < 100 and ext not in ['.txt', '.md', '.csv', '.json']:
            findings.append(MetadataFinding(
                type="consistency",
                description=f"File is unusually small ({size} bytes)",
                severity="low",
                data={"size": size}
            ))
    
    # Check file hashes
    if 'hashes' in metadata and 'md5' in metadata['hashes']:
        md5 = metadata['hashes']['md5']
        
        # Here you could add checks against known malware hashes
        # For example, checking against VirusTotal API or similar
        # This is just a placeholder
        known_malicious_hashes = []  # This would need to be populated
        if md5 in known_malicious_hashes:
            findings.append(MetadataFinding(
                type="security",
                description="File hash matches known malicious file",
                severity="high",
                data={"hash": md5}
            ))
    
    return findings

def redact_metadata(input_path: str, output_path: str, keep_fields: List[str] = None) -> bool:
    """Create a copy of the file with metadata removed."""
    if keep_fields is None:
        keep_fields = []
    
    ext = os.path.splitext(input_path)[1].lower()
    
    try:
        # Create a temporary directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_output = os.path.join(temp_dir, os.path.basename(output_path))
            
            # Handle different file types
            if ext in SUPPORTED_EXTENSIONS['images']:
                # Process image files
                redact_image_metadata(input_path, temp_output, keep_fields)
            elif ext in SUPPORTED_EXTENSIONS['documents']:
                if ext == '.pdf':
                    redact_pdf_metadata(input_path, temp_output, keep_fields)
                else:
                    redact_office_metadata(input_path, temp_output, keep_fields)
            elif ext in SUPPORTED_EXTENSIONS['audio']:
                redact_audio_metadata(input_path, temp_output, keep_fields)
            else:
                # For unsupported types, just copy the file
                shutil.copy2(input_path, temp_output)
                logging.warning(f"Metadata redaction not supported for {ext} files. Created a plain copy.")
            
            # Move from temp location to final destination
            shutil.move(temp_output, output_path)
            
            return True
    except Exception as e:
        logging.error(f"Error during metadata redaction: {e}")
        return False

def redact_image_metadata(input_path: str, output_path: str, keep_fields: List[str]) -> None:
    """Remove metadata from image files."""
    try:
        with Image.open(input_path) as img:
            # Create a new image with the same content but without metadata
            data = list(img.getdata())
            new_img = Image.new(img.mode, img.size)
            new_img.putdata(data)
            
            # Keep only specified metadata if any
            preserved_data = {}
            if keep_fields:
                # Handle specific preservation logic based on image format
                if img.format == 'JPEG' and hasattr(img, '_exif'):
                    exif = img._getexif()
                    if exif:
                        for field in keep_fields:
                            for tag, value in exif.items():
                                tag_name = ExifTags.TAGS.get(tag, str(tag))
                                if tag_name == field or tag == field:
                                    preserved_data[tag] = value
            
            # Save with minimal metadata
            save_kwargs = {}
            if img.format == 'JPEG':
                if preserved_data and 'exif' in keep_fields:
                    try:
                        from PIL import TiffImagePlugin
                        exif_bytes = TiffImagePlugin.ImageFileDirectory_v2()
                        for tag, value in preserved_data.items():
                            exif_bytes[tag] = value
                        save_kwargs['exif'] = exif_bytes.tobytes()
                    except Exception as e:
                        logging.warning(f"Could not preserve EXIF data: {e}")
                
                new_img.save(output_path, format=img.format, **save_kwargs)
            elif img.format == 'PNG':
                # PNG has no standard metadata blocks to preserve
                new_img.save(output_path, format=img.format)
            else:
                new_img.save(output_path, format=img.format)
    
    except Exception as e:
        logging.error(f"Error redacting image metadata: {e}")
        # Fallback: create a clean copy by converting
        try:
            with Image.open(input_path) as img:
                # Convert to a format that strips metadata
                if img.mode == 'RGBA':
                    new_img = Image.new('RGBA', img.size)
                    new_img.paste(img, (0, 0), img)
                else:
                    new_img = Image.new(img.mode, img.size)
                    new_img.paste(img, (0, 0))
                
                new_img.save(output_path)
        except Exception as fallback_error:
            logging.error(f"Fallback image redaction also failed: {fallback_error}")
            # Last resort: just copy the file
            shutil.copy2(input_path, output_path)
            raise RuntimeError("Could not redact image metadata")

def redact_pdf_metadata(input_path: str, output_path: str, keep_fields: List[str]) -> None:
    """Remove metadata from PDF files."""
    try:
        with open(input_path, 'rb') as f_in:
            reader = PdfReader(f_in)
            
            # Create a new PDF writer
            from PyPDF2 import PdfWriter
            writer = PdfWriter()
            
            # Copy all pages
            for page in reader.pages:
                writer.add_page(page)
            
            # Set minimal metadata or preserved fields
            if keep_fields:
                for field in keep_fields:
                    if reader.metadata and field in reader.metadata:
                        writer.add_metadata({f"/{field}": reader.metadata[f"/{field}"]})
            
            # Write the output file
            with open(output_path, 'wb') as f_out:
                writer.write(f_out)
    
    except Exception as e:
        logging.error(f"Error redacting PDF metadata: {e}")
        # Fallback: just copy the file
        shutil.copy2(input_path, output_path)
        raise RuntimeError("Could not redact PDF metadata")

def redact_office_metadata(input_path: str, output_path: str, keep_fields: List[str]) -> None:
    """Remove metadata from Office documents."""
    ext = os.path.splitext(input_path)[1].lower()
    
    try:
        if ext == '.docx' and OPTIONAL_DEPENDENCIES['docx']:
            import docx
            doc = docx.Document(input_path)
            
            # Reset core properties except those in keep_fields
            if not any(field in keep_fields for field in ['author', 'creator']):
                doc.core_properties.author = ''
            if 'title' not in keep_fields:
                doc.core_properties.title = ''
            if 'subject' not in keep_fields:
                doc.core_properties.subject = ''
            if 'comments' not in keep_fields:
                doc.core_properties.comments = ''
            if 'category' not in keep_fields:
                doc.core_properties.category = ''
            if 'keywords' not in keep_fields:
                doc.core_properties.keywords = ''
            
            # Save the document
            doc.save(output_path)
        
        elif ext == '.xlsx' and OPTIONAL_DEPENDENCIES['openpyxl']:
            import openpyxl
            wb = openpyxl.load_workbook(input_path)
            
            # Reset properties except those in keep_fields
            if not any(field in keep_fields for field in ['creator', 'author']):
                wb.properties.creator = ''
            if 'title' not in keep_fields:
                wb.properties.title = ''
            if 'subject' not in keep_fields:
                wb.properties.subject = ''
            if 'description' not in keep_fields:
                wb.properties.description = ''
            if 'category' not in keep_fields:
                wb.properties.category = ''
            if 'keywords' not in keep_fields:
                wb.properties.keywords = ''
            
            # Save the workbook
            wb.save(output_path)
        
        else:
            # For other office formats, we might need additional libraries
            # or external tools like LibreOffice command line
            shutil.copy2(input_path, output_path)
            logging.warning(f"Limited metadata redaction for {ext} files. Some metadata may remain.")
    
    except Exception as e:
        logging.error(f"Error redacting Office document metadata: {e}")
        # Fallback: just copy the file
        shutil.copy2(input_path, output_path)
        raise RuntimeError("Could not redact Office document metadata")

def redact_audio_metadata(input_path: str, output_path: str, keep_fields: List[str]) -> None:
    """Remove metadata from audio files."""
    ext = os.path.splitext(input_path)[1].lower()
    
    try:
        if ext == '.mp3':
            try:
                from mutagen.id3 import ID3, ID3NoHeaderError
                
                # Create a temporary copy first
                shutil.copy2(input_path, output_path)
                
                try:
                    # Try to load ID3 tags
                    audio = ID3(output_path)
                    
                    # If keep_fields is empty, delete all tags
                    if not keep_fields:
                        audio.delete()
                    else:
                        # Otherwise, keep only specified tags
                        tags_to_keep = []
                        for frame in audio.values():
                            frame_name = frame.__class__.__name__
                            if frame_name in keep_fields or any(field.lower() in frame_name.lower() for field in keep_fields):
                                tags_to_keep.append(frame)
                        
                        # Delete all tags
                        audio.delete()
                        
                        # Re-add the ones to keep
                        for frame in tags_to_keep:
                            audio.add(frame)
                    
                    # Save the modified file
                    audio.save()
                
                except ID3NoHeaderError:
                    # No ID3 tags present, nothing to redact
                    pass
            except ImportError:
                logging.warning("Mutagen ID3 not available. Copying file without redaction.")
                shutil.copy2(input_path, output_path)
        
        elif ext in ['.flac', '.ogg']:
            # Handling for FLAC and OGG files
            audio = mutagen.File(input_path)
            
            if audio:
                # Create a copy with audio content
                shutil.copy2(input_path, output_path)
                
                # Load the copy
                new_audio = mutagen.File(output_path)
                
                # If keep_fields is empty, clear all tags
                if not keep_fields:
                    new_audio.clear()
                else:
                    # Keep only specified tags
                    tags_to_remove = []
                    for key in new_audio:
                        if not any(field.lower() in key.lower() for field in keep_fields):
                            tags_to_remove.append(key)
                    
                    for key in tags_to_remove:
                        del new_audio[key]
                
                # Save the modified file
                new_audio.save()
        
        else:
            # For other audio formats, just copy (limited support)
            shutil.copy2(input_path, output_path)
            logging.warning(f"Limited metadata redaction for {ext} files. Some metadata may remain.")
    
    except Exception as e:
        logging.error(f"Error redacting audio metadata: {e}")
        # Fallback: just copy the file
        shutil.copy2(input_path, output_path)
        raise RuntimeError("Could not redact audio metadata")

def compare_fuzzy_hash(file_path: str, reference_hash: str) -> List[MetadataFinding]:
    """Compare file's fuzzy hash with a reference hash to find similarities."""
    findings = []
    
    # Check if ssdeep is available
    if not OPTIONAL_DEPENDENCIES['ssdeep']:
        findings.append(MetadataFinding(
            type="warning",
            description="Fuzzy hash comparison requested but ssdeep is not available",
            severity="low",
            data={"error": "Missing ssdeep dependency"}
        ))
        return findings
    
    try:
        # Compute hash for the current file
        file_hash = ssdeep.hash_from_file(file_path)
        
        # Compare with reference hash
        similarity = ssdeep.compare(file_hash, reference_hash)
        
        # Only report if similarity is significant
        if similarity > 50:
            findings.append(MetadataFinding(
                type="similarity",
                description=f"File has {similarity}% similarity with reference hash",
                severity="medium" if similarity > 80 else "low",
                data={
                    "file_hash": file_hash,
                    "reference_hash": reference_hash,
                    "similarity_score": similarity
                }
            ))
    except Exception as e:
        logging.error(f"Error during fuzzy hash comparison: {e}")
        findings.append(MetadataFinding(
            type="error",
            description="Error during fuzzy hash comparison",
            severity="low",
            data={"error": str(e)}
        ))
    
    return findings

def generate_comparison_html(comparison_results: Dict, file_paths: List[str]) -> str:
    """Generate HTML report for file comparison."""
    file_names = [os.path.basename(path) for path in file_paths]
    
    html_parts = ['<!DOCTYPE html><html><head><title>Metadata Comparison Report</title>',
                 '<style>body{font-family:sans-serif;margin:20px;} .section{margin-bottom:20px;border:1px solid #ddd;padding:15px;border-radius:5px;}',
                 '.different{background-color:#ffe0e0;} table{border-collapse:collapse;width:100%;}',
                 'th,td{text-align:left;padding:8px;border-bottom:1px solid #ddd;} th{background-color:#f2f2f2;}</style></head><body>',
                 '<h1>Metadata Comparison Report</h1>']
    
    # Add generation timestamp
    html_parts.append(f'<p>Generated: {datetime.datetime.now().isoformat()}</p>')
    
    # Add file overview
    html_parts.append('<div class="section"><h2>Files</h2><table><tr><th>File</th><th>Path</th></tr>')
    for i, path in enumerate(file_paths):
        html_parts.append(f'<tr><td>File {i+1}</td><td>{path}</td></tr>')
    html_parts.append('</table></div>')
    
    # Add basic info section
    html_parts.append('<div class="section"><h2>Basic Information</h2><table><tr><th>Property</th>')
    for name in file_names:
        html_parts.append(f'<th>{name}</th>')
    html_parts.append('</tr>')
    
    if 'basic_info' in comparison_results:
        for prop, values in comparison_results['basic_info'].items():
            # Determine if values are different
            unique_values = set(str(v) for v in values if v is not None)
            row_class = ' class="different"' if len(unique_values) > 1 else ''
            
            html_parts.append(f'<tr{row_class}><td>{prop}</td>')
            for value in values:
                html_parts.append(f'<td>{value if value is not None else "N/A"}</td>')
            html_parts.append('</tr>')
    
    html_parts.append('</table></div>')
    
    # Add timestamps section
    if 'timestamps' in comparison_results:
        html_parts.append('<div class="section"><h2>Timestamps</h2><table><tr><th>Timestamp</th>')
        for name in file_names:
            html_parts.append(f'<th>{name}</th>')
        html_parts.append('</tr>')
        
        for timestamp, values in comparison_results['timestamps'].items():
            # Determine if values are different
            unique_values = set(str(v) for v in values if v is not None)
            row_class = ' class="different"' if len(unique_values) > 1 else ''
            
            html_parts.append(f'<tr{row_class}><td>{timestamp}</td>')
            for value in values:
                html_parts.append(f'<td>{value if value is not None else "N/A"}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</table></div>')
    
    # Add hashes section
    if 'hashes' in comparison_results:
        html_parts.append('<div class="section"><h2>File Hashes</h2><table><tr><th>Hash Type</th>')
        for name in file_names:
            html_parts.append(f'<th>{name}</th>')
        html_parts.append('</tr>')
        
        for hash_type, values in comparison_results['hashes'].items():
            # Determine if values are different
            unique_values = set(str(v) for v in values if v is not None)
            row_class = ' class="different"' if len(unique_values) > 1 else ''
            
            html_parts.append(f'<tr{row_class}><td>{hash_type}</td>')
            for value in values:
                html_parts.append(f'<td>{value if value is not None else "N/A"}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</table></div>')
    
    # Add fuzzy hash comparison if available
    if 'fuzzy_hash' in comparison_results:
        html_parts.append('<div class="section"><h2>Fuzzy Hash Comparison</h2>')
        
        if 'error' in comparison_results['fuzzy_hash']:
            html_parts.append(f'<p>Error: {comparison_results["fuzzy_hash"]["error"]}</p>')
        else:
            # Show raw hashes
            html_parts.append('<h3>Fuzzy Hashes</h3><table><tr><th>File</th><th>Fuzzy Hash</th></tr>')
            for i, hash_value in enumerate(comparison_results['fuzzy_hash']['hashes']):
                html_parts.append(f'<tr><td>{file_names[i]}</td><td>{hash_value if hash_value else "N/A"}</td></tr>')
            html_parts.append('</table>')
            
            # Show comparisons
            if 'comparisons' in comparison_results['fuzzy_hash']:
                html_parts.append('<h3>Similarity Analysis</h3><table><tr><th>File 1</th><th>File 2</th><th>Similarity</th></tr>')
                for comp in comparison_results['fuzzy_hash']['comparisons']:
                    similarity = comp['similarity']
                    color = '#ff9999' if similarity < 50 else '#ffcc99' if similarity < 80 else '#99ff99'
                    html_parts.append(f'<tr style="background-color:{color}"><td>{comp["file1"]}</td><td>{comp["file2"]}</td>'
                                    f'<td>{similarity}%</td></tr>')
                html_parts.append('</table>')
        
        html_parts.append('</div>')
    
    # Add metadata differences section
    if 'differences' in comparison_results and comparison_results['differences']:
        html_parts.append('<div class="section"><h2>Key Metadata Differences</h2><table><tr><th>Field</th>')
        for name in file_names:
            html_parts.append(f'<th>{name}</th>')
        html_parts.append('</tr>')
        
        for diff in comparison_results['differences']:
            field = diff['field']
            values = diff['values']
            
            html_parts.append(f'<tr class="different"><td>{field}</td>')
            for value in values:
                html_parts.append(f'<td>{value if value is not None else "N/A"}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</table></div>')
    
    html_parts.append('</body></html>')
    return ''.join(html_parts)

def compare_metadata(results: List[FileMetadata], use_fuzzy_hash: bool = False) -> Dict:
    """Compare metadata between multiple files."""
    comparison = {
        'files': [r.file_path for r in results],
        'basic_info': {},
        'timestamps': {},
        'hashes': {},
        'metadata_fields': {},
        'differences': [],
        'similarities': []
    }
    
    # Compare basic info
    comparison['basic_info']['file_type'] = [r.file_type for r in results]
    comparison['basic_info']['mime_type'] = [r.mime_type for r in results]
    comparison['basic_info']['file_size'] = [r.file_size for r in results]
    
    # Compare timestamps
    comparison['timestamps']['creation_time'] = [r.creation_time for r in results]
    comparison['timestamps']['modification_time'] = [r.modification_time for r in results]
    comparison['timestamps']['access_time'] = [r.access_time for r in results]
    
    # Compare file hashes
    hash_types = set()
    for r in results:
        if r.hashes:
            hash_types.update(r.hashes.keys())
    
    for hash_type in hash_types:
        comparison['hashes'][hash_type] = [r.hashes.get(hash_type) if r.hashes else None for r in results]
    
    # Identify key metadata fields across all files
    all_metadata_fields = set()
    metadata_values = {}
    
    for r in results:
        # Flatten metadata structure
        flat_metadata = {}
        
        def flatten_dict(d, parent_key=''):
            if not isinstance(d, dict):
                return
                
            for k, v in d.items():
                new_key = f"{parent_key}.{k}" if parent_key else k
                if isinstance(v, dict):
                    flatten_dict(v, new_key)
                else:
                    flat_metadata[new_key] = v
        
        flatten_dict(r.metadata)
        all_metadata_fields.update(flat_metadata.keys())
        
        # Store values for comparison
        for field, value in flat_metadata.items():
            if field not in metadata_values:
                metadata_values[field] = [None] * len(results)
            metadata_values[field][results.index(r)] = value
    
    # Add metadata field comparisons
    for field, values in metadata_values.items():
        # Only include fields that exist in multiple files and have different values
        if sum(1 for v in values if v is not None) > 1:
            comparison['metadata_fields'][field] = values
    
    # Identify significant differences
    for field, values in comparison['metadata_fields'].items():
        unique_values = set(str(v) for v in values if v is not None)
        if len(unique_values) > 1:
            comparison['differences'].append({
                'field': field,
                'values': values
            })
    
    # Perform fuzzy hash comparison if requested
    if use_fuzzy_hash and OPTIONAL_DEPENDENCIES['pyssdeep']:
        try:
            import pyssdeep as ssdeep
            
            # Compute fuzzy hashes
            fuzzy_hashes = []
            for r in results:
                try:
                    hash_value = ssdeep.hash_from_file(r.file_path)
                    fuzzy_hashes.append(hash_value)
                except Exception as e:
                    logging.error(f"Error computing fuzzy hash for {r.file_path}: {e}")
                    fuzzy_hashes.append(None)
            
            # Compare each pair
            comparison['fuzzy_hash'] = {
                'hashes': fuzzy_hashes,
                'comparisons': []
            }
            
            for i in range(len(fuzzy_hashes)):
                for j in range(i+1, len(fuzzy_hashes)):
                    if fuzzy_hashes[i] and fuzzy_hashes[j]:
                        similarity = ssdeep.compare(fuzzy_hashes[i], fuzzy_hashes[j])
                        comparison['fuzzy_hash']['comparisons'].append({
                            'file1': os.path.basename(results[i].file_path),
                            'file2': os.path.basename(results[j].file_path),
                            'similarity': similarity
                        })
        except ImportError:
            comparison['fuzzy_hash'] = {
                'error': "ssdeep package not installed"
            }
    
    return comparison

def scan_with_yara_rules(file_path: str, rules_path: str) -> List[MetadataFinding]:
    """Scan a file with YARA rules for additional insights."""
    findings = []
    
    # Check if YARA is available
    if not OPTIONAL_DEPENDENCIES['yara']:
        findings.append(MetadataFinding(
            type="warning",
            description="YARA scanning requested but yara-python is not available",
            severity="low",
            data={"error": "Missing yara-python dependency"}
        ))
        return findings
    
    try:
        # Check if rules_path is a directory or a single file
        if os.path.isdir(rules_path):
            # Compile all .yar files in the directory
            rule_files = [os.path.join(rules_path, f) for f in os.listdir(rules_path) 
                        if f.endswith('.yar') or f.endswith('.yara')]
            if not rule_files:
                logging.warning(f"No YARA rule files found in {rules_path}")
                return []
                
            # Compile all rules
            rules = yara.compile(filepaths={os.path.basename(f): f for f in rule_files})
        else:
            # Compile a single rule file
            rules = yara.compile(filepath=rules_path)
        
        # Match rules against the file
        matches = rules.match(file_path)
        
        # Process matches
        for match in matches:
            # Extract rule metadata if available
            meta = match.meta if hasattr(match, 'meta') else {}
            severity = meta.get('severity', 'medium')
            description = meta.get('description', f"YARA rule match: {match.rule}")
            
            findings.append(MetadataFinding(
                type="yara_match",
                description=description,
                severity=severity,
                data={
                    "rule": match.rule,
                    "tags": match.tags,
                    "meta": meta,
                    "strings": [str(s) for s in match.strings] if hasattr(match, 'strings') else []
                }
            ))
    except Exception as e:
        logging.error(f"Error during YARA scanning: {e}")
        findings.append(MetadataFinding(
            type="error",
            description="Error during YARA scanning",
            severity="low",
            data={"error": str(e)}
        ))
    
    return findings

def search_patterns_in_metadata(metadata: Dict) -> List[MetadataFinding]:
    """Search for patterns like emails, phone numbers in metadata."""
    findings = []
    
    # Flatten metadata for easier searching
    flat_metadata = {}
    
    def flatten_dict(d, parent_key=''):
        for k, v in d.items():
            new_key = f"{parent_key}.{k}" if parent_key else k
            if isinstance(v, dict):
                flatten_dict(v, new_key)
            elif isinstance(v, (str, int, float, bool)):
                flat_metadata[new_key] = str(v)
    
    flatten_dict(metadata)
    
    # Search for patterns in flattened metadata
    for pattern_name, pattern in PRIVACY_CONCERNS.items():
        matches = set()
        for key, value in flat_metadata.items():
            if isinstance(value, str):
                found = pattern.findall(value)
                for match in found:
                    if isinstance(match, tuple):  # Some regex patterns return tuples
                        match = match[0]
                    matches.add(match)
        
        if matches:
            severity = "high" if pattern_name in ('email', 'ssn') else "medium"
            findings.append(MetadataFinding(
                type="privacy",
                description=f"Found {pattern_name} in metadata",
                severity=severity,
                data={"matches": list(matches), "pattern": pattern_name}
            ))
    
    return findings

# --- File Processing Functions ---

def process_files(file_paths: List[str], options: Dict = None) -> List[FileMetadata]:
    """Process multiple files in parallel using a thread pool."""
    if options is None:
        options = {}
    
    results = []
    
    # Determine number of worker threads
    max_workers = options.get('max_workers', min(32, os.cpu_count() + 4))
    
    # Process files in parallel
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_file = {executor.submit(process_file, file_path, options): file_path for file_path in file_paths}
        
        if options.get('show_progress', True) and len(file_paths) > 1:
            # Display progress bar for multiple files
            with tqdm.tqdm(total=len(file_paths), desc="Processing files", unit="file") as pbar:
                for future in concurrent.futures.as_completed(future_to_file):
                    result = future.result()
                    results.append(result)
                    pbar.update(1)
        else:
            # Without progress bar
            for future in concurrent.futures.as_completed(future_to_file):
                results.append(future.result())
    
    return results

def process_file(file_path: str, options: Dict = None) -> FileMetadata:
    """Process a single file and extract its metadata and perform analysis."""
    """Process a single file and extract its metadata."""
    if options is None:
        options = {}
    
    try:
        # Normalize and validate path
        file_path = os.path.abspath(os.path.normpath(file_path))
        if not os.path.isfile(file_path):
            return FileMetadata(
                file_path=file_path,
                file_type="unknown",
                file_size=0,
                mime_type="unknown",
                errors=[f"File not found or is not accessible: {file_path}"]
            )
        
        # Get basic file info
        mime_type, description = detect_file_type(file_path)
        file_size = os.path.getsize(file_path)
        ext = os.path.splitext(file_path)[1].lower()
        
        # Determine file type
        if ext in SUPPORTED_EXTENSIONS['images']:
            file_type = "image"
        elif ext in SUPPORTED_EXTENSIONS['documents']:
            file_type = "document"
        elif ext in SUPPORTED_EXTENSIONS['audio']:
            file_type = "audio"
        elif ext in SUPPORTED_EXTENSIONS['video']:
            file_type = "video"
        elif ext in SUPPORTED_EXTENSIONS['executables']:
            file_type = "executable"
        else:
            file_type = "other"
        
        # Create basic FileMetadata object
        result = FileMetadata(
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            mime_type=mime_type
        )
        
        # Get file hashes if requested
        if not options.get('skip_hashes', False):
            result.hashes = compute_file_hashes(file_path)
        
        # Get timestamps
        timestamps = get_file_timestamps(file_path)
        if 'creation_time' in timestamps:
            result.creation_time = timestamps['creation_time']
        if 'modification_time' in timestamps:
            result.modification_time = timestamps['modification_time']
        if 'access_time' in timestamps:
            result.access_time = timestamps['access_time']
        
        # Extract type-specific metadata
        if file_type == "image" and not options.get('skip_type_specific', False):
            result.metadata = extract_image_metadata(file_path)
        elif file_type == "document" and not options.get('skip_type_specific', False):
            if ext == '.pdf':
                result.metadata = extract_pdf_metadata(file_path)
            else:
                result.metadata = extract_office_document_metadata(file_path)
        elif file_type == "audio" and not options.get('skip_type_specific', False):
            result.metadata = extract_audio_metadata(file_path)
        elif file_type == "video" and not options.get('skip_type_specific', False):
            result.metadata = extract_video_metadata(file_path)
        elif file_type == "executable" and not options.get('skip_type_specific', False):
            result.metadata = extract_executable_metadata(file_path)
        else:
            # Get generic metadata for other file types
            result.metadata = extract_generic_metadata(file_path)
        
        # Analyze metadata if requested
        if not options.get('skip_analysis', False):
            # Run type-specific analysis
            if file_type == "image":
                result.findings.extend(analyze_image_metadata(result.metadata))
            elif file_type == "document":
                result.findings.extend(analyze_document_metadata(result.metadata))
            elif file_type == "audio":
                result.findings.extend(analyze_audio_metadata(result.metadata))
            elif file_type == "video":
                result.findings.extend(analyze_video_metadata(result.metadata))
            elif file_type == "executable":
                result.findings.extend(analyze_executable_metadata(result.metadata))
            
            # Run generic analysis for all file types
            result.findings.extend(analyze_metadata_generic(result.metadata))
            
            # Search for patterns in metadata
            result.findings.extend(search_patterns_in_metadata(result.metadata))
            
            # Custom YARA rule scanning if enabled
            if options.get('yara_rules_path') and os.path.exists(options.get('yara_rules_path')):
                result.findings.extend(scan_with_yara_rules(file_path, options.get('yara_rules_path')))
            
            # Fuzzy hash (ssdeep) comparison if enabled and reference hash is provided
            if options.get('fuzzy_hash_comparison') and options.get('reference_hash'):
                result.findings.extend(compare_fuzzy_hash(file_path, options.get('reference_hash')))
        
        return result
    except Exception as e:
        logging.error(f"Error processing file {file_path}: {e}")
        return FileMetadata(
            file_path=file_path,
            file_type="unknown",
            file_size=os.path.getsize(file_path) if os.path.exists(file_path) else 0,
            mime_type="unknown",
            errors=[f"Processing error: {str(e)}"]
        )

def write_report(results: List[FileMetadata], output_format: str, output_file: str = None) -> None:
    """Write analysis results to file or stdout in specified format."""
    if output_format == 'json':
        # JSON output
        output = json.dumps([r.to_dict() for r in results], indent=2, default=str)
    elif output_format == 'csv':
        # CSV output (simplified, focusing on findings)
        output = "file_path,file_type,finding_type,severity,description\n"
        for result in results:
            base_info = f'"{result.file_path}","{result.file_type}"'
            if result.findings:
                for finding in result.findings:
                    output += f'{base_info},"{finding.type}","{finding.severity}","{finding.description}"\n'
            else:
                output += f'{base_info},"none","none","No findings"\n'
    elif output_format == 'html':
        # HTML report
        html_parts = ['<!DOCTYPE html><html><head><title>Metadata Analysis Report</title>',
                     '<style>body{font-family:sans-serif;margin:20px;} .file{margin-bottom:20px;border:1px solid #ddd;padding:15px;border-radius:5px;}',
                     '.high{color:red;} .medium{color:orange;} .low{color:green;} .finding{margin:10px 0;} table{border-collapse:collapse;width:100%;}',
                     'th,td{text-align:left;padding:8px;border-bottom:1px solid #ddd;} th{background-color:#f2f2f2;}</style></head><body>',
                     '<h1>Metadata Analysis Report</h1>']
        
        # Add generation timestamp
        html_parts.append(f'<p>Generated: {datetime.datetime.now().isoformat()}</p>')
        
        # Add summary section
        html_parts.append('<h2>Summary</h2>')
        html_parts.append('<table><tr><th>File</th><th>Type</th><th>Size</th><th>Findings</th></tr>')
        
        for result in results:
            high_count = sum(1 for f in result.findings if f.severity.lower() == 'high')
            medium_count = sum(1 for f in result.findings if f.severity.lower() == 'medium')
            low_count = sum(1 for f in result.findings if f.severity.lower() == 'low')
            
            findings_cell = f'<span class="high">{high_count} High</span>, ' \
                            f'<span class="medium">{medium_count} Medium</span>, ' \
                            f'<span class="low">{low_count} Low</span>'
            
            file_name = os.path.basename(result.file_path)
            html_parts.append(f'<tr><td>{file_name}</td><td>{result.file_type}</td>'
                             f'<td>{result.file_size:,} bytes</td><td>{findings_cell}</td></tr>')
        
        html_parts.append('</table>')
        
        # Add detailed results for each file
        html_parts.append('<h2>Detailed Results</h2>')
        
        for result in results:
            file_name = os.path.basename(result.file_path)
            html_parts.append(f'<div class="file"><h3>{file_name}</h3>')
            
            # File info
            html_parts.append('<h4>File Information</h4>')
            html_parts.append('<table>')
            html_parts.append(f'<tr><td>Full Path</td><td>{result.file_path}</td></tr>')
            html_parts.append(f'<tr><td>File Type</td><td>{result.file_type}</td></tr>')
            html_parts.append(f'<tr><td>MIME Type</td><td>{result.mime_type}</td></tr>')
            html_parts.append(f'<tr><td>Size</td><td>{result.file_size:,} bytes</td></tr>')
            
            # Add hashes
            if result.hashes:
                for hash_type, hash_value in result.hashes.items():
                    html_parts.append(f'<tr><td>{hash_type.upper()} Hash</td><td>{hash_value}</td></tr>')
            
            # Add timestamps
            if result.creation_time:
                html_parts.append(f'<tr><td>Created</td><td>{result.creation_time}</td></tr>')
            if result.modification_time:
                html_parts.append(f'<tr><td>Modified</td><td>{result.modification_time}</td></tr>')
            
            html_parts.append('</table>')
            
            # Findings
            html_parts.append('<h4>Analysis Findings</h4>')
            
            if result.findings:
                for finding in result.findings:
                    severity_class = finding.severity.lower()
                    html_parts.append(f'<div class="finding {severity_class}">')
                    html_parts.append(f'<strong>[{finding.type.upper()}] {finding.description}</strong>')
                    
                    # Add finding details if present
                    if finding.data:
                        html_parts.append('<ul>')
                        for key, value in finding.data.items():
                            if isinstance(value, dict):
                                html_parts.append(f'<li>{key}:<ul>')
                                for k, v in value.items():
                                    html_parts.append(f'<li>{k}: {v}</li>')
                                html_parts.append('</ul></li>')
                            elif isinstance(value, list):
                                html_parts.append(f'<li>{key}: {", ".join(str(item) for item in value)}</li>')
                            else:
                                html_parts.append(f'<li>{key}: {value}</li>')
                        html_parts.append('</ul>')
                    
                    html_parts.append('</div>')
            else:
                html_parts.append('<p>No notable findings.</p>')
            
            html_parts.append('</div>')
        
        html_parts.append('</body></html>')
        output = ''.join(html_parts)
    else:
        # Default text output
        output_parts = []
        
        for i, result in enumerate(results):
            if i > 0:
                output_parts.append("=" * 70)  # Separator between files
            
            # File info header
            file_name = os.path.basename(result.file_path)
            output_parts.append(f"{Fore.CYAN}File: {file_name}{Style.RESET_ALL}")
            output_parts.append(f"Path: {result.file_path}")
            output_parts.append(f"Type: {result.file_type} ({result.mime_type})")
            output_parts.append(f"Size: {result.file_size:,} bytes")
            
            # Add hashes
            if result.hashes:
                output_parts.append("\nHashes:")
                for hash_type, hash_value in result.hashes.items():
                    output_parts.append(f"  {hash_type.upper()}: {hash_value}")
            
            # Add timestamps
            if result.creation_time or result.modification_time:
                output_parts.append("\nTimestamps:")
                if result.creation_time:
                    output_parts.append(f"  Created: {result.creation_time}")
                if result.modification_time:
                    output_parts.append(f"  Modified: {result.modification_time}")
            
            # Findings
            output_parts.append("\nAnalysis Findings:")
            if result.findings:
                output_parts.append(format_findings(result.findings))
            else:
                output_parts.append("  No notable findings.")
        
        output = "\n".join(output_parts)
    
    # Write to file or stdout
    if output_file:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(output)
        print(f"Report written to {output_file}")
    else:
        print(output)

def format_findings(findings: List[MetadataFinding], format_type: str = 'text') -> str:
    """Format analysis findings for display in various formats."""
    if not findings:
        return "No notable findings."
    
    if format_type == 'json':
        return json.dumps([f.__dict__ for f in findings], indent=2)
    
    # Text/console format with color
    output = []
    severity_colors = {
        'high': Fore.RED,
        'medium': Fore.YELLOW,
        'low': Fore.GREEN,
        'info': Fore.CYAN
    }
    
    for finding in findings:
        severity = finding.severity.lower()
        color = severity_colors.get(severity, '')
        
        # Format finding header
        header = f"[{finding.type.upper()}] {color}{finding.description}{Style.RESET_ALL}"
        output.append(header)
        
        # Format additional data if present
        if finding.data:
            for key, value in finding.data.items():
                if isinstance(value, dict):
                    output.append(f"  {key}:")
                    for k, v in value.items():
                        output.append(f"    {k}: {v}")
                elif isinstance(value, list):
                    output.append(f"  {key}: {', '.join(str(item) for item in value)}")
                else:
                    output.append(f"  {key}: {value}")
        
        output.append("")
    
    return "\n".join(output)

def main():
    """Run the MetaScout CLI Tool."""
    # Create parser with subcommands
    parser = argparse.ArgumentParser(
        description=f"MetaScout v{VERSION} - Advanced File Metadata Analysis Tool",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent(f'''
        Examples:
          metascout analyze image.jpg
          metascout analyze --format json --output report.json document.pdf
          metascout batch --recursive /path/to/directory --output report.html --format html
          metascout compare original.pdf modified.pdf
          
        Supported file types:
          Images: {', '.join(SUPPORTED_EXTENSIONS['images'])}
          Documents: {', '.join(SUPPORTED_EXTENSIONS['documents'])}
          Audio: {', '.join(SUPPORTED_EXTENSIONS['audio'])}
          Video: {', '.join(SUPPORTED_EXTENSIONS['video'])}
          Executables: {', '.join(SUPPORTED_EXTENSIONS['executables'])}
        ''')
    )
    
    # Global options
    parser.add_argument('--verbose', '-v', action='store_true', help='Enable verbose output')
    parser.add_argument('--quiet', '-q', action='store_true', help='Suppress non-essential output')
    parser.add_argument('--log-file', default='metascout.log', help='Log file path')
    parser.add_argument('--version', action='version', version=f'MetaScout v{VERSION}')
    
    # Create subparsers
    subparsers = parser.add_subparsers(dest='command', help='Command to execute')
    
    # 'analyze' command
    analyze_parser = subparsers.add_parser('analyze', help='Analyze a single file')
    analyze_parser.add_argument('file', help='Path to the file to analyze')
    analyze_parser.add_argument('--format', choices=['text', 'json', 'csv', 'html'], default='text',
                               help='Output format (default: text)')
    analyze_parser.add_argument('--output', '-o', help='Output file (default: stdout)')
    analyze_parser.add_argument('--skip-hashes', action='store_true', help='Skip computing file hashes')
    analyze_parser.add_argument('--yara-rules', help='Path to YARA rules file or directory')
    analyze_parser.add_argument('--skip-analysis', action='store_true', help='Skip analysis, extract metadata only')
    
    # 'batch' command
    batch_parser = subparsers.add_parser('batch', help='Process multiple files or directories')
    batch_parser.add_argument('path', help='Path to file or directory to process')
    batch_parser.add_argument('--recursive', '-r', action='store_true', help='Process directories recursively')
    batch_parser.add_argument('--format', choices=['text', 'json', 'csv', 'html'], default='text',
                             help='Output format (default: text)')
    batch_parser.add_argument('--output', '-o', help='Output file (default: stdout)')
    batch_parser.add_argument('--skip-hashes', action='store_true', help='Skip computing file hashes')
    batch_parser.add_argument('--threads', type=int, default=0, 
                             help='Number of worker threads (default: auto)')
    batch_parser.add_argument('--filter', help='Only process files matching this glob pattern')
    batch_parser.add_argument('--exclude', help='Exclude files matching this glob pattern')
    batch_parser.add_argument('--max-files', type=int, default=0, 
                             help='Maximum number of files to process (0 = unlimited)')
    batch_parser.add_argument('--yara-rules', help='Path to YARA rules file or directory')
    
    # 'compare' command
    compare_parser = subparsers.add_parser('compare', help='Compare metadata between files')
    compare_parser.add_argument('files', nargs='+', help='Files to compare')
    compare_parser.add_argument('--format', choices=['text', 'json', 'html'], default='text',
                               help='Output format (default: text)')
    compare_parser.add_argument('--output', '-o', help='Output file (default: stdout)')
    compare_parser.add_argument('--fuzzy-hash', action='store_true', 
                               help='Compare files using fuzzy hashing')
    
    # 'redact' command
    redact_parser = subparsers.add_parser('redact', help='Create a redacted copy with metadata removed')
    redact_parser.add_argument('input_file', help='Input file to redact')
    redact_parser.add_argument('output_file', help='Output file path')
    redact_parser.add_argument('--keep', nargs='+', help='Metadata fields to preserve')
    
    # Parse arguments
    args = parser.parse_args()
    
    # Configure logging
    log_level = logging.DEBUG if args.verbose else logging.INFO
    if args.quiet:
        log_level = logging.WARNING
    
    configure_logging(args.log_file, args.verbose and not args.quiet)
    
    # Show banner unless quiet mode
    if not args.quiet:
        print(f"{Fore.CYAN}MetaScout v{VERSION}{Style.RESET_ALL} - Advanced File Metadata Analysis Tool")
        print(f"{'=' * 60}")
    
    # Process command
    if args.command == 'analyze':
        # Single file analysis
        file_path = os.path.abspath(os.path.normpath(args.file))
        if not os.path.isfile(file_path):
            logging.error(f"Error: '{file_path}' is not a valid file.")
            return 1
        
        # Set up options
        options = {
            'skip_hashes': args.skip_hashes,
            'skip_analysis': args.skip_analysis,
            'yara_rules_path': args.yara_rules
        }
        
        # Process the file
        result = process_file(file_path, options)
        
        # Output results
        write_report([result], args.format, args.output)
    
    elif args.command == 'batch':
        path = os.path.abspath(os.path.normpath(args.path))
        
        if os.path.isfile(path):
            # Single file mode
            files = [path]
        elif os.path.isdir(path):
            # Directory mode
            if args.recursive:
                # Recursive walk
                files = []
                for root, _, filenames in os.walk(path):
                    for filename in filenames:
                        files.append(os.path.join(root, filename))
            else:
                # Non-recursive (just top directory)
                files = [os.path.join(path, f) for f in os.listdir(path) 
                        if os.path.isfile(os.path.join(path, f))]
            
            # Apply file filters if specified
            if args.filter:
                import fnmatch
                files = [f for f in files if fnmatch.fnmatch(os.path.basename(f), args.filter)]
            
            if args.exclude:
                import fnmatch
                files = [f for f in files if not fnmatch.fnmatch(os.path.basename(f), args.exclude)]
            
            # Apply max files limit if specified
            if args.max_files > 0 and len(files) > args.max_files:
                logging.info(f"Limiting to {args.max_files} files out of {len(files)} found")
                files = files[:args.max_files]
        else:
            logging.error(f"Error: '{path}' is not a valid file or directory.")
            return 1
        
        if not files:
            logging.error("No files found to process.")
            return 1
        
        # Set up options
        options = {
            'skip_hashes': args.skip_hashes,
            'max_workers': args.threads if args.threads > 0 else None,
            'show_progress': not args.quiet,
            'yara_rules_path': args.yara_rules
        }
        
        # Process files
        if not args.quiet:
            print(f"Processing {len(files)} files...")
        
        results = process_files(files, options)
        
        # Output results
        write_report(results, args.format, args.output)
    
    elif args.command == 'compare':
        # Compare metadata between files
        if len(args.files) < 2:
            logging.error("Error: At least two files are required for comparison.")
            return 1
        
        # Normalize paths
        file_paths = [os.path.abspath(os.path.normpath(f)) for f in args.files]
        
        # Validate files
        for file_path in file_paths:
            if not os.path.isfile(file_path):
                logging.error(f"Error: '{file_path}' is not a valid file.")
                return 1
        
        # Process files
        options = {'skip_analysis': True}  # Initial metadata extraction only
        results = []
        
        for file_path in file_paths:
            results.append(process_file(file_path, options))
        
        # Perform comparison
        comparison_results = compare_metadata(results, use_fuzzy_hash=args.fuzzy_hash)
        
        # Output comparison
        if args.format == 'json':
            output = json.dumps(comparison_results, indent=2, default=str)
            
            if args.output:
                with open(args.output, 'w', encoding='utf-8') as f:
                    f.write(output)
                print(f"Comparison report written to {args.output}")
            else:
                print(output)
        elif args.format == 'html':
            # Generate HTML comparison report
            html_output = generate_comparison_html(comparison_results, file_paths)
            
            if args.output:
                with open(args.output, 'w', encoding='utf-8') as f:
                    f.write(html_output)
                print(f"Comparison report written to {args.output}")
            else:
                print(html_output)
        else:
            # Text output
            print("File Metadata Comparison Report")
            print("=" * 60)
            
            # Print file info
            for i, file_path in enumerate(file_paths):
                print(f"File {i+1}: {os.path.basename(file_path)}")
                print(f"  Path: {file_path}")
                print(f"  Type: {results[i].file_type}")
                print(f"  Size: {results[i].file_size:,} bytes")
                if results[i].hashes:
                    print(f"  MD5: {results[i].hashes.get('md5', 'N/A')}")
                print()
            
            # Print comparison sections
            for section, details in comparison_results.items():
                if section != 'files':
                    print(f"{section.upper()}")
                    print("-" * 60)
                    
                    if isinstance(details, dict):
                        for key, values in details.items():
                            print(f"  {key}:")
                            for i, value in enumerate(values):
                                value_str = str(value) if value is not None else 'N/A'
                                print(f"    File {i+1}: {value_str}")
                    else:
                        print(f"  {details}")
                    
                    print()
    
    elif args.command == 'redact':
        # Create redacted copy of a file
        input_path = os.path.abspath(os.path.normpath(args.input_file))
        output_path = os.path.abspath(os.path.normpath(args.output_file))
        
        if not os.path.isfile(input_path):
            logging.error(f"Error: Input file '{input_path}' does not exist.")
            return 1
        
        # Get fields to keep
        keep_fields = args.keep if args.keep else []
        
        # Perform redaction
        success = redact_metadata(input_path, output_path, keep_fields)
        
        if success:
            print(f"Created redacted copy at '{output_path}'")
            
            # Analyze the redacted file to confirm
            if not args.quiet:
                print("\nAnalysis of redacted file:")
                result = process_file(output_path)
                write_report([result], 'text')
        else:
            logging.error("Redaction failed. See log for details.")
            return 1
    
    else:
        # No command specified, show help
        parser.print_help()
    
    return 0

# --- Main entry point ---
if __name__ == "__main__":
    try:
        exit_code = main()
        sys.exit(exit_code)
    except KeyboardInterrupt:
        print("\nOperation cancelled by user.")
        sys.exit(1)
    except Exception as e:
        print(f"\nUnexpected error: {e}")
        logging.error(f"Unhandled exception: {e}", exc_info=True)
        sys.exit(1)