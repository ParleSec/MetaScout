#!/usr/bin/env python3
"""
MetaScout Comprehensive Demo Script

This script demonstrates and verifies all major functionality of MetaScout:
- Single file analysis for all supported file types
- Batch processing of multiple files
- Metadata comparison between files
- Metadata redaction

The script creates realistic sample files with actual metadata for thorough testing.

Usage:
    python metascout_demo.py [--keep-temp-files]
"""

import os
import sys
import shutil
import argparse
import tempfile
import subprocess
import urllib.request
import io
import platform
import datetime
import base64
import json
import re
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple, Union

# Try to import additional helpful libraries
try:
    from tqdm import tqdm
    HAS_TQDM = True
except ImportError:
    HAS_TQDM = False
    print("Warning: tqdm not found, progress bars will be disabled")

# ANSI colors for pretty output (Windows-compatible)
class Colors:
    HEADER = '\033[95m'
    BLUE = '\033[94m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

# Enable colored output on Windows
if platform.system() == 'Windows':
    try:
        import colorama
        colorama.init()
    except ImportError:
        pass

# Output directories
TEMP_DIR = Path(tempfile.mkdtemp(prefix="metascout_demo_"))
OUTPUT_DIR = TEMP_DIR / "output"
REDACTED_DIR = TEMP_DIR / "redacted"
SAMPLE_FILES_DIR = TEMP_DIR / "sample_files"

# File type constants
FILE_TYPES = {
    'text': ['.txt', '.md', '.csv'],
    'document': ['.pdf', '.docx', '.doc', '.odt'],
    'image': ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.webp'],
    'audio': ['.mp3', '.wav', '.flac', '.ogg'],
    'video': ['.mp4', '.avi', '.mkv', '.mov'],
    'archive': ['.zip', '.tar.gz'],
    'executable': ['.exe', '.dll', '.so'],
}

# Check for MetaScout availability
try:
    # First try to import from source for development mode
    import metascout
    from metascout import process_file, process_files
    from metascout.core.models import FileMetadata, MetadataFinding
    from metascout.core.utils import detect_file_type, compute_file_hashes
    from metascout.operations.analyze import write_report
    from metascout.operations.batch import process_directory, filter_files
    from metascout.operations.compare import compare_metadata
    from metascout.operations.redact import redact_metadata
    HAS_METASCOUT = True
except ImportError:
    HAS_METASCOUT = False
    print("Warning: MetaScout not found in path. Demo will create files but not analyze them.")

#-------------------------------------------------------------------------
# Helper functions
#-------------------------------------------------------------------------
def print_header(text):
    """Print a formatted header."""
    print(f"\n{Colors.HEADER}{Colors.BOLD}{'=' * 80}{Colors.ENDC}")
    print(f"{Colors.HEADER}{Colors.BOLD} {text} {Colors.ENDC}")
    print(f"{Colors.HEADER}{Colors.BOLD}{'=' * 80}{Colors.ENDC}\n")

def print_subheader(text):
    """Print a formatted subheader."""
    print(f"\n{Colors.BLUE}{Colors.BOLD}---- {text} ----{Colors.ENDC}\n")

def print_success(text):
    """Print a success message."""
    print(f"{Colors.GREEN}✓ {text}{Colors.ENDC}")

def print_warning(text):
    """Print a warning message."""
    print(f"{Colors.YELLOW}⚠ {text}{Colors.ENDC}")

def print_error(text):
    """Print an error message."""
    print(f"{Colors.RED}✗ {text}{Colors.ENDC}")

def check_dependency(module_name: str, package_name: Optional[str] = None) -> bool:
    """Check if a dependency is available and provide install info if not."""
    try:
        __import__(module_name)
        return True
    except ImportError:
        pkg = package_name or module_name
        print_warning(f"{module_name} not installed. Some functionality may be limited.")
        print(f"  Install with: pip install {pkg}")
        return False

def download_file(url: str, output_path: Path) -> bool:
    """Download a file from a URL."""
    try:
        # Create progress bar if tqdm is available
        if HAS_TQDM:
            with urllib.request.urlopen(url) as response:
                file_size = int(response.info().get('Content-Length', 0))
                desc = f"Downloading {output_path.name}"
                
                with tqdm(total=file_size, unit='B', unit_scale=True, desc=desc) as pbar:
                    with open(output_path, 'wb') as out_file:
                        chunk_size = 1024
                        while True:
                            chunk = response.read(chunk_size)
                            if not chunk:
                                break
                            out_file.write(chunk)
                            pbar.update(len(chunk))
        else:
            urllib.request.urlretrieve(url, output_path)
            
        print_success(f"Downloaded {output_path.name}")
        return True
    except Exception as e:
        print_error(f"Failed to download file: {e}")
        return False

def create_directory_structure():
    """Create the necessary directory structure for the demo."""
    # Create main directories
    SAMPLE_FILES_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)
    REDACTED_DIR.mkdir(exist_ok=True)
    
    # Create subdirectories by file type
    for file_type in FILE_TYPES:
        (SAMPLE_FILES_DIR / file_type).mkdir(exist_ok=True)
    
    print_success(f"Created directory structure at {TEMP_DIR}")

def convert_file_size(size_in_bytes: int) -> str:
    """Convert bytes to a human-readable string."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_in_bytes < 1024 or unit == 'GB':
            return f"{size_in_bytes:.2f} {unit}"
        size_in_bytes /= 1024
    return f"{size_in_bytes:.2f} GB"  # Fallback, should never reach here

#-------------------------------------------------------------------------
# Sample file creation functions
#-------------------------------------------------------------------------
def create_text_files():
    """Create sample text files with various forms of metadata."""
    print_subheader("Creating text files")
    text_dir = SAMPLE_FILES_DIR / "text"
    
    # Create a basic text file with metadata in its content
    basic_txt = text_dir / "basic.txt"
    with open(basic_txt, "w", encoding="utf-8") as f:
        f.write("SIMPLE TEXT FILE WITH METADATA\n")
        f.write("==============================\n\n")
        f.write(f"Created: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("Author: MetaScout Demo\n")
        f.write("Organization: Example Corp.\n\n")
        f.write("This file contains some sample metadata in its content.\n")
        f.write("Email: demo@example.com\n")
        f.write("Phone: 555-123-4567\n")
        f.write("SSN: 123-45-6789\n")
        f.write("Credit Card: 4111-1111-1111-1111\n\n")
        f.write("Website: https://example.com/metadata\n")
    print_success(f"Created text file with PII: {basic_txt}")

    # Create a markdown file with front matter metadata
    markdown_file = text_dir / "document.md"
    with open(markdown_file, "w", encoding="utf-8") as f:
        f.write("---\n")
        f.write("title: Sample Markdown Document\n")
        f.write("author: MetaScout Demo\n")
        f.write(f"date: {datetime.datetime.now().strftime('%Y-%m-%d')}\n")
        f.write("tags: [metadata, demo, example]\n")
        f.write("---\n\n")
        f.write("# Sample Markdown Document\n\n")
        f.write("This is a sample markdown document with YAML front matter metadata.\n\n")
        f.write("## Features\n\n")
        f.write("- Front matter metadata\n")
        f.write("- Markdown formatting\n")
        f.write("- Sample content\n\n")
        f.write(f"Created on {datetime.datetime.now().isoformat()}\n")
    print_success(f"Created markdown file with front matter: {markdown_file}")

    # Create a CSV file with metadata
    csv_file = text_dir / "data.csv"
    with open(csv_file, "w", encoding="utf-8") as f:
        f.write("# Dataset: Sample Customer Data\n")
        f.write("# Author: MetaScout Demo\n")
        f.write(f"# Created: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("# Source: Example Corp.\n")
        f.write("id,name,email,phone,address\n")
        f.write("1,John Doe,john.doe@example.com,555-123-4567,\"123 Main St, Anytown, CA 12345\"\n")
        f.write("2,Jane Smith,jane.smith@example.com,555-987-6543,\"456 Oak Ave, Somewhere, NY 67890\"\n")
        f.write("3,Bob Johnson,bob.johnson@example.com,555-567-8901,\"789 Pine Rd, Nowhere, TX 10111\"\n")
        f.write("4,Alice Brown,alice.brown@example.com,555-234-5678,\"321 Elm Blvd, Everywhere, FL 21314\"\n")
    print_success(f"Created CSV file with header comments: {csv_file}")

    # Create a simple HTML file
    html_file = text_dir / "webpage.html"
    with open(html_file, "w", encoding="utf-8") as f:
        f.write("<!DOCTYPE html>\n")
        f.write("<html>\n")
        f.write("<head>\n")
        f.write("    <title>Sample HTML Document</title>\n")
        f.write("    <meta name=\"author\" content=\"MetaScout Demo\">\n")
        f.write("    <meta name=\"description\" content=\"A sample HTML file for testing MetaScout\">\n")
        f.write("    <meta name=\"keywords\" content=\"metadata,demo,test,metascout\">\n")
        f.write("    <meta name=\"generator\" content=\"MetaScout Demo Script\">\n")
        f.write("    <meta name=\"created\" content=\"" + datetime.datetime.now().isoformat() + "\">\n")
        f.write("    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">\n")
        f.write("    <meta name=\"copyright\" content=\"Example Corp. 2025\">\n")
        f.write("    <meta name=\"robots\" content=\"index, follow\">\n")
        f.write("    <meta property=\"og:title\" content=\"Sample HTML Document\">\n")
        f.write("    <meta property=\"og:description\" content=\"A sample HTML file for testing MetaScout\">\n")
        f.write("    <meta property=\"og:image\" content=\"https://example.com/image.jpg\">\n")
        f.write("    <meta property=\"og:url\" content=\"https://example.com/sample\">\n")
        f.write("</head>\n")
        f.write("<body>\n")
        f.write("    <h1>Sample HTML Document</h1>\n")
        f.write("    <p>This is a sample HTML document for testing MetaScout.</p>\n")
        f.write("    <p>It contains metadata in its meta tags.</p>\n")
        f.write("    <p>Contact: <a href=\"mailto:demo@example.com\">demo@example.com</a></p>\n")
        f.write("    <p>Phone: <a href=\"tel:555-123-4567\">555-123-4567</a></p>\n")
        f.write("    <div class=\"footer\">\n")
        f.write("        <p>&copy; 2025 Example Corp. All rights reserved.</p>\n")
        f.write("    </div>\n")
        f.write("    <!-- Hidden comment with metadata: Created by MetaScout Demo on " + datetime.datetime.now().isoformat() + " -->\n")
        f.write("</body>\n")
        f.write("</html>\n")
    print_success(f"Created HTML file with meta tags: {html_file}")

    return [basic_txt, markdown_file, csv_file, html_file]

def create_document_files():
    """Create sample document files (PDF, DOCX, etc.)."""
    print_subheader("Creating document files")
    doc_dir = SAMPLE_FILES_DIR / "document"
    
    # Check for dependencies
    has_reportlab = check_dependency("reportlab")
    has_docx = check_dependency("docx", "python-docx")
    
    files_created = []
    
    # Create a PDF file with metadata
    if has_reportlab:
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            pdf_file = doc_dir / "report.pdf"
            c = canvas.Canvas(str(pdf_file), pagesize=letter)
            
            # Add metadata
            c.setTitle("Sample PDF Document")
            c.setAuthor("MetaScout Demo")
            c.setSubject("Testing MetaScout")
            c.setKeywords("metadata,demo,test,metascout")
            c.setCreator("MetaScout Demo Script")
            
            # Add content
            c.drawString(100, 750, "Sample PDF Document")
            c.drawString(100, 735, "Created for MetaScout testing")
            c.drawString(100, 720, f"Generated on {datetime.datetime.now().isoformat()}")
            c.drawString(100, 705, "This document has metadata that can be extracted with MetaScout")
            c.drawString(100, 690, "Contact: demo@example.com")
            c.drawString(100, 675, "Phone: 555-123-4567")
            
            # Add some example sensitive information
            c.drawString(100, 640, "EXAMPLE SENSITIVE DATA (for testing only):")
            c.drawString(120, 625, "SSN: 123-45-6789")
            c.drawString(120, 610, "Credit Card: 4111-1111-1111-1111")
            c.drawString(120, 595, "Password: p@ssw0rd123")
            
            c.save()
            print_success(f"Created PDF file with metadata: {pdf_file}")
            files_created.append(pdf_file)
        except Exception as e:
            print_error(f"Failed to create PDF file: {e}")
    else:
        # Download a sample PDF if reportlab is not available
        pdf_url = "https://www.africau.edu/images/default/sample.pdf"  # Generic sample PDF
        pdf_file = doc_dir / "sample.pdf"
        if download_file(pdf_url, pdf_file):
            files_created.append(pdf_file)
    
    # Create a DOCX file with metadata
    if has_docx:
        try:
            from docx import Document
            from docx.shared import Inches
            
            docx_file = doc_dir / "document.docx"
            document = Document()
            
            # Add metadata
            document.core_properties.author = "MetaScout Demo"
            document.core_properties.title = "Sample Word Document"
            document.core_properties.subject = "Testing MetaScout"
            document.core_properties.keywords = "metadata,demo,test,metascout"
            document.core_properties.category = "Demo"
            document.core_properties.comments = "This is a sample document for testing metadata extraction"
            
            # Add content
            document.add_heading('Sample Word Document', 0)
            
            p = document.add_paragraph('This is a sample Word document with ')
            p.add_run('metadata').bold = True
            p.add_run(' that can be extracted with MetaScout.')
            
            document.add_heading('Contact Information', level=1)
            document.add_paragraph('Email: demo@example.com')
            document.add_paragraph('Phone: 555-123-4567')
            
            document.add_heading('Example Sensitive Data (for testing only)', level=1)
            document.add_paragraph('SSN: 123-45-6789')
            document.add_paragraph('Credit Card: 4111-1111-1111-1111')
            document.add_paragraph('Password: p@ssw0rd123')
            
            document.save(docx_file)
            print_success(f"Created DOCX file with metadata: {docx_file}")
            files_created.append(docx_file)
        except Exception as e:
            print_error(f"Failed to create DOCX file: {e}")
    else:
        # Download a sample DOCX if python-docx is not available
        docx_url = "https://file-examples.com/wp-content/uploads/2017/02/file-sample_100kB.docx"
        docx_file = doc_dir / "sample.docx"
        if download_file(docx_url, docx_file):
            files_created.append(docx_file)
    
    return files_created

def create_image_files():
    """Create sample image files with EXIF and other metadata."""
    print_subheader("Creating image files")
    image_dir = SAMPLE_FILES_DIR / "image"
    
    # Check for dependencies
    has_pil = check_dependency("PIL", "pillow")
    has_piexif = check_dependency("piexif")
    
    files_created = []
    
    # Create a JPEG file with EXIF metadata
    if has_pil and has_piexif:
        try:
            from PIL import Image
            import piexif
            
            # Create a simple image
            jpg_file = image_dir / "photo.jpg"
            img = Image.new('RGB', (800, 600), color=(73, 109, 137))
            
            # Create EXIF data
            exif_dict = {
                "0th": {
                    piexif.ImageIFD.Make: "MetaScout".encode(),
                    piexif.ImageIFD.Model: "Demo Camera".encode(),
                    piexif.ImageIFD.Software: "MetaScout Demo".encode(),
                    piexif.ImageIFD.Copyright: "Example Corp. 2025".encode()
                },
                "Exif": {
                    piexif.ExifIFD.DateTimeOriginal: datetime.datetime.now().strftime("%Y:%m:%d %H:%M:%S").encode(),
                    piexif.ExifIFD.LensMake: "Demo Lens".encode(),
                    piexif.ExifIFD.UserComment: "This is a test image with EXIF data".encode()
                },
                "GPS": {
                    piexif.GPSIFD.GPSLatitudeRef: "N".encode(),
                    piexif.GPSIFD.GPSLatitude: ((40, 1), (44, 1), (0, 1)),
                    piexif.GPSIFD.GPSLongitudeRef: "W".encode(),
                    piexif.GPSIFD.GPSLongitude: ((73, 1), (59, 1), (0, 1)),
                }
            }
            
            exif_bytes = piexif.dump(exif_dict)
            
            # Draw something on the image
            if hasattr(img, 'draw'):
                from PIL import ImageDraw, ImageFont
                draw = ImageDraw.Draw(img)
                draw.rectangle([50, 50, 750, 550], outline=(255, 255, 255), width=2)
                draw.text((400, 300), "MetaScout Demo Image", fill=(255, 255, 255), anchor="mm")
            
            # Save with EXIF data
            img.save(jpg_file, "JPEG", exif=exif_bytes, quality=95)
            print_success(f"Created JPEG file with EXIF data: {jpg_file}")
            files_created.append(jpg_file)
            
            # Create a PNG file with metadata
            png_file = image_dir / "graphic.png"
            img = Image.new('RGBA', (800, 600), color=(73, 109, 137, 255))
            
            # Draw something on the image
            if hasattr(img, 'draw'):
                from PIL import ImageDraw, ImageFont
                draw = ImageDraw.Draw(img)
                draw.rectangle([50, 50, 750, 550], outline=(255, 255, 255, 255), width=2)
                draw.text((400, 300), "MetaScout Demo PNG", fill=(255, 255, 255, 255), anchor="mm")
            
            # Add metadata using PngInfo
            meta = Image.PngImagePlugin.PngInfo()
            meta.add_text("Author", "MetaScout Demo")
            meta.add_text("Title", "Sample PNG Image")
            meta.add_text("Description", "This is a sample PNG image with metadata")
            meta.add_text("Copyright", "Example Corp. 2025")
            meta.add_text("Creation Time", datetime.datetime.now().isoformat())
            meta.add_text("Software", "MetaScout Demo Script")
            
            # Save with metadata
            img.save(png_file, "PNG", pnginfo=meta)
            print_success(f"Created PNG file with metadata: {png_file}")
            files_created.append(png_file)
            
        except Exception as e:
            print_error(f"Failed to create image file with metadata: {e}")
    
    # If PIL is not available or creation failed, download sample images
    if not has_pil or len(files_created) == 0:
        # Download a sample JPEG with EXIF data
        jpg_url = "https://raw.githubusercontent.com/ianare/exif-samples/master/jpg/gps/DSCN0010.jpg"
        jpg_file = image_dir / "sample.jpg"
        if download_file(jpg_url, jpg_file):
            files_created.append(jpg_file)
        
        # Download a sample PNG
        png_url = "https://upload.wikimedia.org/wikipedia/commons/4/47/PNG_transparency_demonstration_1.png"
        png_file = image_dir / "sample.png"
        if download_file(png_url, png_file):
            files_created.append(png_file)
    
    return files_created

def create_audio_files():
    """Create sample audio files with ID3 tags and other metadata."""
    print_subheader("Creating audio files")
    audio_dir = SAMPLE_FILES_DIR / "audio"
    
    # Check for dependencies
    has_mutagen = check_dependency("mutagen")
    
    files_created = []
    
    # Try to create a valid MP3 file with ID3 tags
    if has_mutagen:
        try:
            import mutagen
            from mutagen.id3 import ID3, TIT2, TPE1, TALB, TDRC, TCON, COMM, WXXX
            
            # First try to download a small valid MP3 file
            mp3_url = "https://file-examples.com/wp-content/uploads/2017/11/file_example_MP3_700KB.mp3"
            mp3_file = audio_dir / "song.mp3"
            
            if download_file(mp3_url, mp3_file):
                # Add ID3 tags
                try:
                    # Try to load existing tags
                    tags = ID3(mp3_file)
                except:
                    # Create new tags if none exist
                    tags = ID3()
                
                # Set tags
                tags.add(TIT2(encoding=3, text="Sample Song"))
                tags.add(TPE1(encoding=3, text="MetaScout Demo"))
                tags.add(TALB(encoding=3, text="Test Album"))
                tags.add(TDRC(encoding=3, text=str(datetime.datetime.now().year)))
                tags.add(TCON(encoding=3, text="Test"))
                tags.add(COMM(encoding=3, lang="eng", desc="Comment", text="This is a sample file for MetaScout testing"))
                tags.add(WXXX(encoding=3, url="https://example.com/metadata"))
                
                # Save tags
                tags.save(mp3_file)
                print_success(f"Added ID3 tags to MP3 file: {mp3_file}")
                files_created.append(mp3_file)
                
                # Try to create a WAV file if we have the MP3
                wav_file = audio_dir / "audio.wav"
                try:
                    # Try to convert MP3 to WAV if pydub is available
                    try:
                        from pydub import AudioSegment
                        audio = AudioSegment.from_mp3(mp3_file)
                        audio.export(wav_file, format="wav")
                        print_success(f"Created WAV file: {wav_file}")
                        files_created.append(wav_file)
                    except ImportError:
                        # Just download a sample WAV
                        wav_url = "https://file-examples.com/wp-content/uploads/2017/11/file_example_WAV_1MG.wav"
                        if download_file(wav_url, wav_file):
                            files_created.append(wav_file)
                except Exception as e:
                    print_error(f"Failed to create WAV file: {e}")
            else:
                print_error("Failed to download MP3 file")
                
                # Fallback: create a tiny MP3-like file with ID3 tags
                mp3_file = audio_dir / "dummy.mp3"
                with open(mp3_file, "wb") as f:
                    # ID3v2 header + empty MPEG frame
                    f.write(b"ID3\x03\x00\x00\x00\x00\x00\x06TPE1\x00\x00\x00\x0D\x00\x00\x03MetaScout Demo")
                    f.write(b"\xFF\xFB\x90\x44\x00" + b"\x00" * 512)  # Fake MPEG frame
                
                print_warning(f"Created dummy MP3-like file: {mp3_file}")
                files_created.append(mp3_file)
        except Exception as e:
            print_error(f"Failed to create audio file with metadata: {e}")
    
    # If mutagen is not available or creation failed, download sample audio files
    if not has_mutagen or len(files_created) == 0:
        # Download a sample MP3
        mp3_url = "https://file-examples.com/wp-content/uploads/2017/11/file_example_MP3_700KB.mp3"
        mp3_file = audio_dir / "sample.mp3"
        if download_file(mp3_url, mp3_file):
            files_created.append(mp3_file)
        
        # Download a sample WAV
        wav_url = "https://file-examples.com/wp-content/uploads/2017/11/file_example_WAV_1MG.wav"
        wav_file = audio_dir / "sample.wav"
        if download_file(wav_url, wav_file):
            files_created.append(wav_file)
    
    return files_created

def create_video_files():
    """Create or download sample video files with metadata."""
    print_subheader("Creating video files")
    video_dir = SAMPLE_FILES_DIR / "video"
    
    files_created = []
    
    # Download a sample MP4 file
    mp4_url = "https://file-examples.com/wp-content/uploads/2017/04/file_example_MP4_480_1_5MG.mp4"
    mp4_file = video_dir / "sample.mp4"
    if download_file(mp4_url, mp4_file):
        files_created.append(mp4_file)
    
    return files_created

def create_archive_files():
    """Create sample archive files with included files and metadata."""
    print_subheader("Creating archive files")
    archive_dir = SAMPLE_FILES_DIR / "archive"
    
    files_created = []
    
    # Create a ZIP file with some content
    try:
        import zipfile
        
        zip_file = archive_dir / "archive.zip"
        with zipfile.ZipFile(zip_file, "w") as zf:
            # Create internal directory structure
            # Add a text file to the zip
            info = zipfile.ZipInfo("readme.txt")
            info.date_time = datetime.datetime.now().timetuple()[:6]
            info.compress_type = zipfile.ZIP_DEFLATED
            info.comment = b"Created by MetaScout Demo Script"
            zf.writestr(info, "This is a sample file inside a ZIP archive for MetaScout testing.\n"
                             "It contains some sample metadata for extraction tests.\n\n"
                             "Author: MetaScout Demo\n"
                             "Email: demo@example.com\n"
                             "Phone: 555-123-4567\n"
                             "SSN: 123-45-6789\n"
                             "Credit Card: 4111-1111-1111-1111\n")
            
            # Add a simple HTML file
            info = zipfile.ZipInfo("page.html")
            info.date_time = datetime.datetime.now().timetuple()[:6]
            info.compress_type = zipfile.ZIP_DEFLATED
            zf.writestr(info, """<!DOCTYPE html>
<html>
<head>
    <title>Sample HTML in ZIP</title>
    <meta name="author" content="MetaScout Demo">
    <meta name="description" content="A sample HTML file in a ZIP archive">
</head>
<body>
    <h1>Sample HTML in ZIP</h1>
    <p>This HTML file is stored inside a ZIP archive.</p>
</body>
</html>""")
            
            # Add a simple JSON file with metadata
            info = zipfile.ZipInfo("data.json")
            info.date_time = datetime.datetime.now().timetuple()[:6]
            info.compress_type = zipfile.ZIP_DEFLATED
            json_data = {
                "metadata": {
                    "title": "Sample JSON Data",
                    "author": "MetaScout Demo",
                    "created": datetime.datetime.now().isoformat(),
                    "description": "A sample JSON file in a ZIP archive"
                },
                "data": [
                    {"id": 1, "name": "John Doe", "email": "john.doe@example.com"},
                    {"id": 2, "name": "Jane Smith", "email": "jane.smith@example.com"},
                    {"id": 3, "name": "Bob Johnson", "email": "bob.johnson@example.com"}
                ]
            }
            zf.writestr(info, json.dumps(json_data, indent=2))
        
        print_success(f"Created ZIP file with metadata: {zip_file}")
        files_created.append(zip_file)
    except Exception as e:
        print_error(f"Failed to create ZIP file: {e}")
    
    return files_created

def create_executable_files():
    """Create sample executable files with metadata."""
    print_subheader("Creating executable files")
    executable_dir = SAMPLE_FILES_DIR / "executable"
    
    files_created = []
    
    # Try to create a valid PE file
    try:
        # Option 1: On Windows, copy a small system executable
        if platform.system() == 'Windows':
            try:
                system32 = os.path.join(os.environ['SystemRoot'], 'System32')
                source_exe = os.path.join(system32, 'notepad.exe')  # Small Windows executable
                exe_file = executable_dir / "windows_app.exe"
                
                if os.path.exists(source_exe):
                    shutil.copy2(source_exe, exe_file)
                    print_success(f"Created Windows EXE from system file: {exe_file}")
                    files_created.append(exe_file)
                else:
                    raise FileNotFoundError("Could not find notepad.exe")
            except Exception as e:
                print_error(f"Failed to copy system executable: {e}")
        
        # Option 2: Create a minimal valid PE file
        exe_file = executable_dir / "sample.exe"
        with open(exe_file, 'wb') as f:
            # DOS Header
            f.write(b'MZ')                     # Magic number
            f.write(b'\x00' * 58)              # DOS stub
            f.write(b'\x40\x00\x00\x00')       # PE header offset at 0x40
            
            # PE Header
            f.write(b'PE\x00\x00')             # PE Signature
            f.write(b'\x4c\x01')               # Machine (x86)
            f.write(b'\x01\x00')               # Number of sections
            
            # Timestamp (current time in seconds since epoch)
            timestamp = int(datetime.datetime.now().timestamp())
            f.write(timestamp.to_bytes(4, byteorder='little'))
            
            f.write(b'\x00' * 8)               # Other header fields
            f.write(b'\xe0\x00\x02\x01')       # Characteristics
            f.write(b'\x0b\x01')               # Magic (for optional header)
            
            # PeInfo (metadata)
            company_name = b"Example Corp."
            product_name = b"MetaScout Demo"
            file_description = b"Sample executable for MetaScout testing"
            original_filename = b"sample.exe"
            
            # Write the metadata in a format similar to PE resources
            # (This is simplified and not a real PE resource section)
            f.write(b'\x00ResourceInfo\x00')
            f.write(company_name + b'\x00')
            f.write(product_name + b'\x00')
            f.write(file_description + b'\x00')
            f.write(original_filename + b'\x00')
            f.write(b'V1.0.0' + b'\x00')
            
            # Pad to a reasonable file size
            f.write(b'\x00' * 1024)
        
        print_success(f"Created minimal PE file: {exe_file}")
        files_created.append(exe_file)
        
        # Option 3: if on Linux, create a minimal ELF file
        if platform.system() == 'Linux':
            elf_file = executable_dir / "linux_app"
            try:
                # Try to copy a small Linux executable
                source_elf = "/bin/ls"  # Small Linux binary
                if os.path.exists(source_elf):
                    shutil.copy2(source_elf, elf_file)
                    print_success(f"Created Linux ELF from system file: {elf_file}")
                    files_created.append(elf_file)
            except Exception as e:
                print_error(f"Failed to create ELF file: {e}")
    except Exception as e:
        print_error(f"Failed to create executable file: {e}")
        
        # Fallback: create a very basic PE-like file
        exe_file = executable_dir / "minimal.exe"
        with open(exe_file, "wb") as f:
            # Just write the MZ header
            f.write(b"MZ")
            f.write(b"\x00" * 64)
        
        print_warning(f"Created minimal PE-like file: {exe_file}")
        files_created.append(exe_file)
    
    return files_created

def create_sample_files():
    """Create all sample files for testing."""
    print_header("Creating Sample Files")
    
    # Create directory structure
    create_directory_structure()
    
    # Create different types of files
    all_files = []
    all_files.extend(create_text_files())
    all_files.extend(create_document_files())
    all_files.extend(create_image_files())
    all_files.extend(create_audio_files())
    all_files.extend(create_video_files())
    all_files.extend(create_archive_files())
    all_files.extend(create_executable_files())
    
    print(f"\nCreated {len(all_files)} sample files across {len(FILE_TYPES)} file types")
    
    return all_files

#-------------------------------------------------------------------------
# Test functions
#-------------------------------------------------------------------------
def test_single_file_analysis(sample_files):
    """Test analyzing each sample file individually."""
    print_header("Testing Single File Analysis")
    
    if not HAS_METASCOUT:
        print_warning("MetaScout not available. Skipping analysis tests.")
        return False
    
    # Analyze a subset of files to avoid making the demo too long
    test_files = []
    # Select one file of each main type
    file_types = list(FILE_TYPES.keys())
    
    for file_type in file_types:
        type_dir = SAMPLE_FILES_DIR / file_type
        if type_dir.exists():
            files = list(type_dir.glob("*"))
            if files:
                test_files.append(files[0])
    
    # If no files were found, use all sample files
    if not test_files:
        test_files = sample_files[:5]  # Limit to 5 files
    
    analysis_results = []
    
    for file_path in test_files:
        print_subheader(f"Analyzing {file_path.relative_to(SAMPLE_FILES_DIR)}")
        
        try:
            # Process the file
            result = process_file(file_path)
            analysis_results.append(result)
            
            # Print basic info
            print(f"File type: {result.file_type}")
            print(f"MIME type: {result.mime_type}")
            print(f"File size: {result.file_size} bytes ({convert_file_size(result.file_size)})")
            
            # Print findings count
            finding_count = len(result.findings)
            if finding_count > 0:
                print_success(f"Found {finding_count} metadata findings")
                
                # Print findings by severity
                high_count = sum(1 for f in result.findings if f.severity.lower() == 'high')
                medium_count = sum(1 for f in result.findings if f.severity.lower() == 'medium')
                low_count = sum(1 for f in result.findings if f.severity.lower() == 'low')
                
                if high_count > 0:
                    print(f"  {Colors.RED}{high_count} High{Colors.ENDC} severity findings")
                if medium_count > 0:
                    print(f"  {Colors.YELLOW}{medium_count} Medium{Colors.ENDC} severity findings")
                if low_count > 0:
                    print(f"  {Colors.GREEN}{low_count} Low{Colors.ENDC} severity findings")
                
                # Print first few findings
                print("\nTop findings:")
                for i, finding in enumerate(sorted(result.findings, 
                                                 key=lambda f: 0 if f.severity.lower() == 'high' else 
                                                               1 if f.severity.lower() == 'medium' else 2)[:3]):
                    severity_color = Colors.RED if finding.severity.lower() == 'high' else \
                                    Colors.YELLOW if finding.severity.lower() == 'medium' else \
                                    Colors.GREEN
                    print(f"  {i+1}. [{severity_color}{finding.severity.upper()}{Colors.ENDC}] {finding.type}: {finding.description}")
                
                if finding_count > 3:
                    print(f"  ... and {finding_count - 3} more findings")
            else:
                print_warning("No metadata findings detected")
            
            # Generate a report file
            OUTPUT_DIR.mkdir(exist_ok=True)
            output_file = OUTPUT_DIR / f"{file_path.stem}_report.txt"
            
            # Write report
            write_report([result], "text", str(output_file))
            
            if output_file.exists():
                print_success(f"Report saved to {output_file.relative_to(TEMP_DIR)}")
            else:
                print_error(f"Failed to create report file")
            
        except Exception as e:
            print_error(f"Error analyzing {file_path.name}: {str(e)}")
    
    if analysis_results:
        # Generate a combined report in HTML format
        combined_output = OUTPUT_DIR / "analysis_report.html"
        write_report(analysis_results, "html", str(combined_output))
        print_success(f"Combined HTML report saved to {combined_output.relative_to(TEMP_DIR)}")
    
    return True

def test_batch_processing(sample_files_dir):
    """Test batch processing of multiple files."""
    print_header("Testing Batch Processing")
    
    if not HAS_METASCOUT:
        print_warning("MetaScout not available. Skipping batch processing tests.")
        return False
    
    try:
        # Process each file type directory separately
        for file_type in FILE_TYPES:
            type_dir = sample_files_dir / file_type
            if not type_dir.exists() or not any(type_dir.iterdir()):
                continue
                
            print_subheader(f"Batch processing {file_type} files")
            
            # Create output file
            OUTPUT_DIR.mkdir(exist_ok=True)
            output_file = OUTPUT_DIR / f"{file_type}_batch_report.html"
            
            # Use the process_directory function
            try:
                results = process_directory(
                    str(type_dir),
                    recursive=False,
                    file_filter=None,
                    exclude_filter=None
                )
                
                # Print results summary
                print_success(f"Processed {len(results)} {file_type} files")
                
                # Count findings by severity
                high_count = sum(sum(1 for f in r.findings if f.severity.lower() == 'high') for r in results)
                medium_count = sum(sum(1 for f in r.findings if f.severity.lower() == 'medium') for r in results)
                low_count = sum(sum(1 for f in r.findings if f.severity.lower() == 'low') for r in results)
                
                print(f"Found {high_count + medium_count + low_count} total findings:")
                if high_count > 0:
                    print(f"  {Colors.RED}{high_count} High{Colors.ENDC} severity findings")
                if medium_count > 0:
                    print(f"  {Colors.YELLOW}{medium_count} Medium{Colors.ENDC} severity findings")
                if low_count > 0:
                    print(f"  {Colors.GREEN}{low_count} Low{Colors.ENDC} severity findings")
                
                # Generate HTML report
                write_report(results, "html", str(output_file))
                
                if output_file.exists():
                    print_success(f"Batch report saved to {output_file.relative_to(TEMP_DIR)}")
                else:
                    print_error("Failed to create batch report file")
            except Exception as e:
                print_error(f"Error processing directory {type_dir}: {e}")
        
        # Test filtering
        print_subheader("Testing file filtering")
        
        # Filter by extension
        all_files = [str(f) for f in sample_files_dir.glob("**/*")]
        
        # Filter image files
        img_extensions = tuple(FILE_TYPES['image'])
        img_files = filter_files(all_files, include_pattern="*.*", exclude_pattern=None)
        img_files = [f for f in img_files if f.lower().endswith(img_extensions)]
        
        print_success(f"Filtered {len(img_files)} image files")
        
        # Filter executable files
        exe_extensions = tuple(FILE_TYPES['executable'])
        exe_files = filter_files(all_files, include_pattern="*.*", exclude_pattern=None)
        exe_files = [f for f in exe_files if f.lower().endswith(exe_extensions)]
        
        print_success(f"Filtered {len(exe_files)} executable files")
        
        return True
    except Exception as e:
        print_error(f"Error in batch processing: {str(e)}")
        return False

def test_comparison(sample_files):
    """Test comparing metadata between files."""
    print_header("Testing Metadata Comparison")
    
    if not HAS_METASCOUT:
        print_warning("MetaScout not available. Skipping comparison tests.")
        return False
    
    # Find some interesting files to compare
    text_files = list(SAMPLE_FILES_DIR.glob("text/*.txt"))
    image_files = list(SAMPLE_FILES_DIR.glob("image/*.jpg"))
    
    if not text_files or not image_files:
        # Fallback to any two files
        if len(sample_files) < 2:
            print_error("Need at least two files for comparison testing")
            return False
        file1, file2 = sample_files[:2]
    else:
        # Compare a text file and an image file
        file1, file2 = text_files[0], image_files[0]
    
    print_subheader(f"Comparing {file1.relative_to(SAMPLE_FILES_DIR)} and {file2.relative_to(SAMPLE_FILES_DIR)}")
    
    try:
        # Process both files
        result1 = process_file(file1)
        result2 = process_file(file2)
        
        # Compare metadata
        comparison = compare_metadata([result1, result2])
        
        # Print some comparison results
        print(f"Basic info comparison:")
        for key, values in comparison.get('basic_info', {}).items():
            print(f"  {key}: {values}")
        
        # Print differences
        diff_count = len(comparison.get('differences', []))
        if diff_count > 0:
            print(f"\nDifferences found: {diff_count}")
            for i, diff in enumerate(comparison.get('differences', [])[:3]):
                print(f"  {i+1}. {diff.get('field')}: {diff.get('values')}")
            
            if diff_count > 3:
                print(f"  ... and {diff_count - 3} more differences")
        else:
            print("\nNo significant differences found")
        
        # Generate comparison report
        OUTPUT_DIR.mkdir(exist_ok=True)
        output_file = OUTPUT_DIR / "comparison_report.json"
        
        # Write comparison to file (simple JSON dump since we may not have the formatted report function)
        with open(output_file, "w") as f:
            json.dump(comparison, f, indent=2, default=str)
        
        if output_file.exists():
            print_success(f"Comparison report saved to {output_file.relative_to(TEMP_DIR)}")
        else:
            print_error("Failed to create comparison report file")
        
        return True
    except Exception as e:
        print_error(f"Error in metadata comparison: {str(e)}")
        return False

def test_redaction(sample_files):
    """Test redacting metadata from files."""
    print_header("Testing Metadata Redaction")
    
    if not HAS_METASCOUT:
        print_warning("MetaScout not available. Skipping redaction tests.")
        return False
    
    REDACTED_DIR.mkdir(exist_ok=True)
    
    # Select one file of each main type for redaction
    redaction_test_files = []
    for file_type in FILE_TYPES:
        type_dir = SAMPLE_FILES_DIR / file_type
        if type_dir.exists():
            files = list(type_dir.glob("*"))
            if files:
                redaction_test_files.append(files[0])
    
    # If no files were found, use all sample files
    if not redaction_test_files:
        redaction_test_files = sample_files[:3]  # Limit to 3 files
    
    for file_path in redaction_test_files:
        print_subheader(f"Redacting metadata from {file_path.relative_to(SAMPLE_FILES_DIR)}")
        
        try:
            # Create output path
            output_path = REDACTED_DIR / f"redacted_{file_path.name}"
            
            # Redact metadata
            success = redact_metadata(
                str(file_path),
                str(output_path),
                keep_fields=[]  # Redact all metadata
            )
            
            if success and output_path.exists():
                print_success(f"Redacted file saved to {output_path.relative_to(TEMP_DIR)}")
                
                # Verify redaction by analyzing the redacted file
                original = process_file(file_path)
                redacted = process_file(output_path)
                
                original_count = len(original.findings)
                redacted_count = len(redacted.findings)
                
                if redacted_count < original_count:
                    print_success(f"Verified redaction: {original_count} findings reduced to {redacted_count}")
                    
                    # Show what was removed
                    if original_count > 0:
                        print("\nRemoved findings:")
                        for finding in original.findings:
                            # Check if this finding is not in redacted
                            if not any(rf.description == finding.description for rf in redacted.findings):
                                severity_color = Colors.RED if finding.severity.lower() == 'high' else \
                                                Colors.YELLOW if finding.severity.lower() == 'medium' else \
                                                Colors.GREEN
                                print(f"  - [{severity_color}{finding.severity.upper()}{Colors.ENDC}] {finding.description}")
                else:
                    print_warning(f"Redaction may not be complete: {original_count} findings in original, {redacted_count} in redacted")
                    
                    # Show what remains
                    if redacted_count > 0:
                        print("\nRemaining findings in redacted file:")
                        for finding in redacted.findings:
                            severity_color = Colors.RED if finding.severity.lower() == 'high' else \
                                            Colors.YELLOW if finding.severity.lower() == 'medium' else \
                                            Colors.GREEN
                            print(f"  - [{severity_color}{finding.severity.upper()}{Colors.ENDC}] {finding.description}")
            else:
                print_error(f"Failed to redact {file_path.name}")
                
        except Exception as e:
            print_error(f"Error redacting {file_path.name}: {str(e)}")
    
    # Test selective redaction
    if len(redaction_test_files) > 0:
        print_subheader("Testing selective redaction (keeping specific metadata)")
        
        try:
            file_path = redaction_test_files[0]
            output_path = REDACTED_DIR / f"selective_{file_path.name}"
            
            # Specific fields to keep
            keep_fields = ['title', 'author', 'creator']
            
            # Redact metadata but keep specified fields
            success = redact_metadata(
                str(file_path),
                str(output_path),
                keep_fields=keep_fields
            )
            
            if success and output_path.exists():
                print_success(f"Selectively redacted file saved to {output_path.relative_to(TEMP_DIR)}")
                print(f"Fields preserved: {', '.join(keep_fields)}")
            else:
                print_error(f"Failed to selectively redact {file_path.name}")
        except Exception as e:
            print_error(f"Error in selective redaction: {str(e)}")
    
    return True

def show_summary():
    """Show a summary of all operations tested."""
    print_header("MetaScout Demo Summary")
    
    print(f"Components tested:")
    if HAS_METASCOUT:
        try:
            from metascout.extractors import EXTRACTORS
            from metascout.analyzers import ANALYZERS
            
            print(f"- {len(EXTRACTORS)} extractors")
            print(f"- {len(ANALYZERS)} analyzers")
            print(f"- Single file analysis")
            print(f"- Batch processing")
            print(f"- Metadata comparison")
            print(f"- Metadata redaction")
        except ImportError:
            print("- MetaScout components (details unavailable)")
    else:
        print("- Sample file creation only (MetaScout not available)")
    
    # Count generated files
    sample_count = sum(1 for _ in SAMPLE_FILES_DIR.glob('**/*'))
    output_count = sum(1 for _ in OUTPUT_DIR.glob('*')) if OUTPUT_DIR.exists() else 0
    redacted_count = sum(1 for _ in REDACTED_DIR.glob('*')) if REDACTED_DIR.exists() else 0
    
    print(f"\nFiles generated:")
    print(f"- {sample_count} sample files with metadata")
    print(f"- {output_count} analysis reports")
    print(f"- {redacted_count} redacted files")
    
    print(f"\nOutput files available in:")
    print(f"- Sample files: {SAMPLE_FILES_DIR}")
    print(f"- Analysis reports: {OUTPUT_DIR}")
    print(f"- Redacted files: {REDACTED_DIR}")
    
    print(f"\nMetaScout demo completed successfully!")

def clean_up(keep_temp_files=False):
    """Clean up temporary files."""
    if keep_temp_files:
        print(f"\nTemporary files preserved at: {TEMP_DIR}")
    else:
        shutil.rmtree(TEMP_DIR)
        print("\nTemporary files removed")

#-------------------------------------------------------------------------
# Main function
#-------------------------------------------------------------------------
def main():
    """Run the MetaScout demo."""
    parser = argparse.ArgumentParser(description="MetaScout Demo Script")
    parser.add_argument("--keep-temp-files", action="store_true", 
                        help="Keep temporary files after running")
    parser.add_argument("--skip-tests", action="store_true",
                        help="Only create sample files, skip tests")
    args = parser.parse_args()
    
    print_header("MetaScout Comprehensive Demo")
    print(f"Starting demo at: {datetime.datetime.now().isoformat()}")
    print(f"Temporary directory: {TEMP_DIR}")
    
    try:
        # Check if MetaScout is available
        if not HAS_METASCOUT:
            print_warning("MetaScout not available. Demo will create sample files but not perform analysis.")
        
        # Create sample files
        sample_files = create_sample_files()
        
        # Run tests unless skipped
        if not args.skip_tests and HAS_METASCOUT:
            test_single_file_analysis(sample_files)
            test_batch_processing(SAMPLE_FILES_DIR)
            test_comparison(sample_files)
            test_redaction(sample_files)
        
        # Show summary
        show_summary()
        
    except Exception as e:
        print_error(f"Unhandled error: {str(e)}")
        import traceback
        traceback.print_exc()
    finally:
        # Clean up
        clean_up(args.keep_temp_files)

if __name__ == "__main__":
    main()